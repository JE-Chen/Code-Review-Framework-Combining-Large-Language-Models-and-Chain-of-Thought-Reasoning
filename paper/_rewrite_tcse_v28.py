"""Produce TCSE_v2.8.docx from TCSE_v2.7.docx (original untouched): the framework
base model becomes gemma-4-31B-it in prose, and the Qwen name moves INTO 表一's
"Ours" column header so the quantitative results stay attributed to the
configuration that produced them. Prose outside tables no longer names Qwen; the
gemma replay status (no judge scores yet — no gemma numbers cited) is recorded
in §6.2. One-off helper, mirrors the underscore-prefixed paper/ scripts.
Usage: .venv/Scripts/python.exe paper/_rewrite_tcse_v28.py"""
import copy
import re
import sys

from docx import Document
from docx.oxml.ns import qn

sys.stdout.reconfigure(encoding="utf-8")
SRC, DST = "paper/TCSE_v2.7.docx", "paper/TCSE_v2.8.docx"
doc = Document(SRC)
body = doc.element.body
W_T, W_P, W_R, W_RPR = qn("w:t"), qn("w:p"), qn("w:r"), qn("w:rPr")


def paras():
    return body.findall(".//" + W_P)


def para_text(p):
    return "".join(t.text or "" for t in p.findall(".//" + W_T))


def set_text(t, s):
    t.text = s
    if s != s.strip():
        t.set(qn("xml:space"), "preserve")


def _rewrite_spans(spans, idx, end, new):
    """Blank out [idx, end) across the runs, writing new into the first hit."""
    first = True
    for a, b, t in spans:
        if b <= idx or a >= end:
            continue
        seg = t.text or ""
        lo, hi = max(idx - a, 0), min(end - a, len(seg))
        set_text(t, seg[:lo] + (new if first else "") + seg[hi:])
        first = False


def replace_once_in_para(p, old, new):
    ts = p.findall(".//" + W_T)
    joined, spans = "", []
    for t in ts:
        s = t.text or ""
        spans.append((len(joined), len(joined) + len(s), t))
        joined += s
    idx = joined.find(old)
    if idx == -1:
        return False
    _rewrite_spans(spans, idx, idx + len(old), new)
    return True


def replace_everywhere(old, new):
    if old in new:
        raise SystemExit(f"old 為 new 之子字串，會無窮替換：{old[:40]}…")
    n = 0
    for p in paras():
        while old in para_text(p):
            before = para_text(p)
            replace_once_in_para(p, old, new)
            if para_text(p) == before:
                raise SystemExit(f"替換無進展：{old[:40]}…")
            n += 1
    if n == 0:
        raise SystemExit(f"定點改寫未命中：{old[:40]}…")
    return n


# ---------- 1. 表格外散文去 Qwen 化（基底改述為 gemma-4-31B-it） ----------
EDITS = [
    # §1.2 構想
    ("轉移至輕量化學生模型（Qwen3-Coder-30B-A3B-Instruct）",
     "轉移至輕量化學生模型（現行基底為 gemma-4-31B-it）"),
    # §3.1 知識蒸餾與微調管線
    ("以該批資料對學生模型（Qwen3-Coder-30B-A3B-Instruct）進行監督式微調",
     "以該批資料對學生模型（gemma-4-31B-it）進行監督式微調"),
    ("微調採 QLoRA 技術：模型權重先以 4-bit 量化以大幅降低顯示卡記憶體佔用，"
     "再僅針對少量低秩適配器（LoRA adapter）參數進行更新，使 30B 等級之模型亦可"
     "於有限硬體下高效訓練。此處有一關鍵取捨：訓練以 4-bit 量化省記憶體，推論卻"
     "改以 bf16 進行，因 A3B 之混合專家（MoE）結構於新版函式庫會在 4-bit 下展開"
     "（densify）而致記憶體溢位。訓練後 LoRA 適配器保持未與基座合併並以 CPU "
     "暫存，以保留動態切換之彈性。",
     "微調採低秩適應：僅針對少量低秩適配器（LoRA adapter）參數進行更新，使 30B "
     "等級之模型亦可於有限硬體下高效訓練；視部署硬體之記憶體條件，模型權重可先以"
     " 4-bit 量化載入（QLoRA，本文表一至表三之實驗配置採此）或直接以 bf16 載入"
     "訓練（現行基底採此）。訓練後 LoRA 適配器保持未與基座合併，以保留動態切換之"
     "彈性。"),
    # §3.2 輸入層
    ("啟動時載入審查模型 Qwen3-Coder-30B（Alibaba 釋出之 30B 程式碼基礎模型）"
     "與訓練完成之 LoRA 權重",
     "啟動時載入審查模型 gemma-4-31B-it（Google 釋出之開放權重模型）與訓練完成之"
     " LoRA 權重"),
    # §3.2 推論層
    ("送入經 LoRA 微調之 Qwen3-Coder-30B 模型執行",
     "送入經 LoRA 微調之 gemma-4-31B-it 模型執行"),
    # §4.2 評估客觀性（去除括弧之模型名）
    ("由獨立於本框架審查模型（Qwen3-Coder-30B）之外部大型語言模型（如 GPT-5）擔任",
     "由獨立於本框架審查模型之外部大型語言模型（如 GPT-5）擔任"),
    # §6.2 限制與未來工作：補跨模型重放既成事實
    ("此等統計穩健性與模組級消融之檢驗，連同上列各項，均留待後續逐項驗證。",
     "此等統計穩健性與模組級消融之檢驗，連同上列各項，均留待後續逐項驗證；另，"
     "跨模型重放之第一步已完成——框架基底現已更換為 gemma-4-31B-it 並加掛專屬 "
     "LoRA 轉接器，於同一 44 筆基準資料與同一管線端到端重放並歸檔逐案輸出，本文"
     "表一至表三之數字均屬前代學生模型配置（如表一欄名所列）之結果，gemma 系列之"
     "評審評分尚未執行，本文不引用其分數，其量化比較仍屬未來工作。"),
]
for old, new in EDITS:
    print(f"改寫 {replace_everywhere(old, new)} 處：{new[:28]}…")

# ---------- 2. 表一「Ours」欄名補模型名（表內，保留結果歸因） ----------
table1 = doc.tables[0]
cell = next(
    (c for c in table1.rows[0].cells if c.text.strip() == "Ours"), None
)
if cell is None:
    raise SystemExit("表一找不到 Ours 欄名儲存格")
first_p = cell.paragraphs[0]._p
extra_p = copy.deepcopy(first_p)
wts = extra_p.findall(".//" + W_T)
if not wts:
    raise SystemExit("表一 Ours 欄名儲存格缺文字節點")
set_text(wts[0], "Qwen3-Coder-30B")
for t in wts[1:]:
    set_text(t, "")
first_p.addnext(extra_p)
print("表一：Ours 欄名補入 Qwen3-Coder-30B")

# ---------- 3. 全形標點正規化（1:1 字元映射，按原 run 長度回填） ----------
PUNCT_MAP = {",": "，", ";": "，", ":": "：", "?": "？", "!": "！"}


def _cjkish(ch):
    return ("一" <= ch <= "鿿" or "　" <= ch <= "〿"
            or "＀" <= ch <= "￯" or ch in "—…•")


def _nearest_visible(seq):
    """First non-space char in seq, or '' when none."""
    return next((c for c in seq if c != " "), "")


def _paren_is_cjk(text, j, i):
    """True when the (j, i) paren pair sits in CJK context."""
    seg_cjk = any(_cjkish(c) for c in text[j + 1:i])
    prev = _nearest_visible(reversed(text[:j]))
    nxt = _nearest_visible(text[i + 1:])
    prev_cjk = bool(prev) and _cjkish(prev)
    nxt_cjk = bool(nxt) and _cjkish(nxt)
    return seg_cjk or prev_cjk or nxt_cjk


def _convert_parens(text, chars):
    """Full-width-ify balanced ASCII parens that sit in CJK context."""
    stack = []
    for i, ch in enumerate(chars):
        if ch == "(":
            stack.append(i)
        elif ch == ")" and stack:
            j = stack.pop()
            if _paren_is_cjk(text, j, i):
                chars[j], chars[i] = "（", "）"


def _convert_marks(chars):
    """Full-width-ify mapped punctuation adjacent to CJK text."""
    for i, ch in enumerate(chars):
        if ch in PUNCT_MAP:
            prev = _nearest_visible(reversed(chars[:i]))
            nxt = _nearest_visible(chars[i + 1:])
            if _cjkish(prev) or _cjkish(nxt):
                chars[i] = PUNCT_MAP[ch]


def convert_punct(text):
    chars = list(text)
    _convert_parens(text, chars)
    _convert_marks(chars)
    return "".join(chars)


np = 0
for p in paras():
    ts = p.findall(".//" + W_T)
    joined = "".join(t.text or "" for t in ts)
    conv = convert_punct(joined)
    if conv != joined:
        pos = 0
        for t in ts:
            n = len(t.text or "")
            set_text(t, conv[pos:pos + n])
            pos += n
        np += 1
print(f"標點正規化：{np} 段")

# ---------- 4. 字型：每個含文字之 run 顯式四 slot ----------
# 中英文與全形標點 run 一律顯式設字型；純符號 run（✓/emoji/數學）保留原字型
HAS_CONTENT = re.compile(r"[A-Za-z0-9一-鿿　-〿＀-￯]")
nf = 0
for r in body.findall(".//" + W_R):
    s = "".join(t.text or "" for t in r.findall(W_T))
    if not s or not HAS_CONTENT.search(s):
        continue
    rpr = r.find(W_RPR)
    if rpr is None:
        rpr = r.makeelement(W_RPR, {})
        r.insert(0, rpr)
    rf = rpr.find(qn("w:rFonts"))
    if rf is None:
        rf = rpr.makeelement(qn("w:rFonts"), {})
        rpr.insert(0, rf)
    for slot in ("w:ascii", "w:hAnsi", "w:cs"):
        rf.set(qn(slot), "Times New Roman")
    rf.set(qn("w:eastAsia"), "標楷體")
    nf += 1
print(f"字型設定：{nf} 個 run")

doc.save(DST)
print(f"已輸出 {DST}")
