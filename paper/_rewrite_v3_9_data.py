"""Apply the DATA-DEPENDENT v3.9 revisions (objective re-experiment results).

Reads 論文_v3.9.before-data.docx (the data-independent v3.9, a snapshot of the
earlier batch) and writes the final 論文_v3.9.docx. The original v3.8 is never
touched. Every number here is traceable to
datas/Results/2026-07-19-gemma4-experiment/ (verified: coverage.json /
precision.json / ground_truth -> recall 197/333=0.592, 166/333=0.498,
precision 474/480 & 329/333 =0.988, F1 0.740/0.663, sev-weighted 0.602/0.510,
per-case 26/14/4; our_score_claude / our_score_gpt56sol -> Δ table & means;
code_to_detect -> LOC 3241/73.7/65/37/150, GT 333/7.6/1/17, sev 7/60/181/85,
7≡8 duplicate).

Change set (spec: paper/_revision_v3.9_data_spec.md D1–D7):
  D1  delete old 表七–表十一 (unsourced Qwen 5-dim GPT-5/Gemini); add objective
      bug-detection table + LLM-judge prose table; keep+relabel human tables as
      prior-gen Qwen; renumber tables/TOC/cross-refs.
  D2  rewrite §5.3.2 as objective bug-detection main result + honest dual-metric.
  D3  §5.3.3/RQ1 -> descriptive CRSCORE++ comparison (drop 優於/全面提升/所以呢).
  D4  reframe RQ2/RQ3 (§1.3) around issue-detection recall.
  D5  three-tier conclusion (§6.1/§6.2), drop +34/+2/三者缺一不可/可用→可信賴.
  D6  §4.1 dataset-statistics table + 7≡8 disclosure; §6.3 limitations
      (synthetic data, model-derived ground truth, 43 unique, single benchmark,
      RAG degenerate 0/44, fine-tune un-ablatable on deployed model).
  D7  paper_rule throughout (full-width punct, no 破折號/分號, no AI-tell, 繁體).

Fonts are normalised at the end (four-slot rFonts, skipping math/symbol runs),
so a separate _fix_fonts pass is not required. Usage:
  .venv/Scripts/python.exe paper/_rewrite_v3_9_data.py
"""
import copy
import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from docx.table import Table
from docx.text.paragraph import Paragraph

sys.stdout.reconfigure(encoding="utf-8")
BASE = Path(__file__).parent
SRC = BASE / "論文_v3.9.before-data.docx"
DST = BASE / "論文_v3.9.docx"
DEPLOYMENT_PNG = BASE.parent / "datas" / "Architecture" / "v3.9" / "工程部署架構.png"


# --------------------------------------------------------------------------- #
# low-level helpers
# --------------------------------------------------------------------------- #
def set_run_fonts(run_el):
    rpr = run_el.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    for slot in ("w:ascii", "w:hAnsi", "w:cs"):
        rfonts.set(qn(slot), "Times New Roman")
    rfonts.set(qn("w:eastAsia"), "標楷體")


def is_tof(p):
    return "table of figures" in (p.style.name or "").lower()


def find_body_para(doc, anchor, exact=False):
    hits = []
    for p in doc.paragraphs:
        if is_tof(p):
            continue
        t = p.text.strip()
        if (t == anchor) if exact else (anchor in p.text):
            hits.append(p)
    if len(hits) != 1:
        raise SystemExit(f"BODY ANCHOR not unique ({len(hits)}): {anchor[:44]}")
    return hits[0]


def rebuild(p, new_text):
    """Collapse a paragraph to a single run carrying the original rPr plus the
    four explicit rFonts slots (safe for plain prose / caption paragraphs)."""
    src_rpr = None
    for r in p.runs:
        if r.text.strip():
            src_rpr = r._element.find(qn("w:rPr"))
            break
    for r in list(p.runs):
        r._element.getparent().remove(r._element)
    run = p.add_run(new_text)
    if src_rpr is not None:
        run._element.insert(0, copy.deepcopy(src_rpr))
    set_run_fonts(run._element)


def rewrite_paragraph(doc, label, anchor, new_text):
    rebuild(find_body_para(doc, anchor), new_text)
    print(f"OK rewrite {label}")


def delete_paragraph(doc, label, anchor):
    p = find_body_para(doc, anchor)
    p._p.getparent().remove(p._p)
    print(f"OK delete paragraph {label}")


def edit_substr(doc, label, anchor, old, new):
    p = find_body_para(doc, anchor)
    txt = p.text
    if txt.count(old) != 1:
        raise SystemExit(f"OLD not unique ({txt.count(old)}) [{label}]: {old[:40]}")
    rebuild(p, txt.replace(old, new))
    print(f"OK edit {label}")


# --------------------------------------------------------------------------- #
# table + paragraph construction
# --------------------------------------------------------------------------- #
def _set_cell(cell, text):
    for extra in cell.paragraphs[1:]:
        extra._p.getparent().remove(extra._p)
    p = cell.paragraphs[0]
    for r in list(p.runs):
        r._element.getparent().remove(r._element)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_run_fonts(run._element)
    run.font.size = Pt(14)


def _resize_rows(tbl, n_target):
    while len(tbl.rows) > n_target:
        tr = tbl.rows[-1]._tr
        tr.getparent().remove(tr)
    while len(tbl.rows) < n_target:
        last_tr = tbl.rows[-1]._tr
        new_tr = copy.deepcopy(last_tr)
        last_tr.addnext(new_tr)


def make_para_like(after_el, parent, fmt, text):
    ppr, rpr = fmt
    new_p = OxmlElement("w:p")
    after_el.addnext(new_p)
    if ppr is not None:
        new_p.append(copy.deepcopy(ppr))
    para = Paragraph(new_p, parent)
    run = para.add_run(text)
    if rpr is not None:
        run._element.insert(0, copy.deepcopy(rpr))
    set_run_fonts(run._element)
    return new_p


def para_fmt(p):
    ppr = p._p.find(qn("w:pPr"))
    rpr = None
    for r in p.runs:
        if r.text.strip():
            rpr = r._element.find(qn("w:rPr"))
            break
    return (copy.deepcopy(ppr) if ppr is not None else None,
            copy.deepcopy(rpr) if rpr is not None else None)


def insert_table_block(after_el, parent, tmpl_tbl, cap_fmt, desc_fmt,
                       caption, rows, desc):
    cap_p = make_para_like(after_el, parent, cap_fmt, caption)
    new_tbl = copy.deepcopy(tmpl_tbl)
    cap_p.addnext(new_tbl)
    tbl = Table(new_tbl, parent)
    _resize_rows(tbl, len(rows))
    for i, vals in enumerate(rows):
        cells = tbl.rows[i].cells
        for j, v in enumerate(vals):
            _set_cell(cells[j], v)
    desc_p = make_para_like(new_tbl, parent, desc_fmt, desc)
    print(f"OK insert table: {caption[:16]}")
    return desc_p


def insert_figure_block(after_para, parent, caption_fmt, desc_fmt,
                        image_path, caption, desc):
    image_p = OxmlElement("w:p")
    after_para._p.addnext(image_p)
    para = Paragraph(image_p, parent)
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run()
    run.add_picture(str(image_path), width=Inches(6.8))
    cap_p = make_para_like(image_p, parent, caption_fmt, caption)
    desc_p = make_para_like(cap_p, parent, desc_fmt, desc)
    print(f"OK insert figure: {caption}")
    return desc_p


def delete_table_block(doc, caption_exact):
    cap = find_body_para(doc, caption_exact, exact=True)
    cap_el = cap._p
    tbl_el = cap_el.getnext()
    if tbl_el is None or tbl_el.tag != qn("w:tbl"):
        raise SystemExit(f"no table after caption: {caption_exact}")
    desc_el = tbl_el.getnext()
    if desc_el is None or desc_el.tag != qn("w:p"):
        raise SystemExit(f"no desc after table: {caption_exact}")
    par = cap_el.getparent()
    par.remove(desc_el)
    par.remove(tbl_el)
    par.remove(cap_el)
    print(f"OK delete block: {caption_exact[:18]}")


def rename_caption(doc, old_exact, new):
    rebuild(find_body_para(doc, old_exact, exact=True), new)
    print(f"OK caption: {old_exact[:10]} -> {new[:10]}")


def replace_table_after_caption(doc, parent, caption_exact, rows):
    cap = find_body_para(doc, caption_exact, exact=True)
    tbl = Table(cap._p.getnext(), parent)
    _resize_rows(tbl, len(rows))
    for i, values in enumerate(rows):
        for j, value in enumerate(values):
            _set_cell(tbl.rows[i].cells[j], value)
    print(f"OK replace table: {caption_exact[:18]}")


def replace_table_header_after_caption(doc, parent, caption_exact, headers):
    cap = find_body_para(doc, caption_exact, exact=True)
    tbl = Table(cap._p.getnext(), parent)
    for j, value in enumerate(headers):
        _set_cell(tbl.rows[0].cells[j], value)
    print(f"OK replace table header: {caption_exact[:18]}")


# --------------------------------------------------------------------------- #
# TOC (表次目錄) helpers
# --------------------------------------------------------------------------- #
def _tof_caption(p):
    """Caption portion of a table-of-tables entry (text before the tab+page)."""
    return p.text.split("\t")[0].strip()


def _tof_all(doc):
    return [
        p for p in doc.paragraphs
        if is_tof(p) and _tof_caption(p).startswith(("表", "圖"))
    ]


def find_tof(doc, cur):
    hits = [p for p in _tof_all(doc) if _tof_caption(p) == cur]
    if len(hits) != 1:
        raise SystemExit(f"TOF not unique ({len(hits)}): {cur}")
    return hits[0]


def _set_tof_caption(p_el, new_caption):
    ts = list(p_el.iter(qn("w:t")))
    if len(ts) < 2:
        raise SystemExit(f"TOF too few w:t: {new_caption}")
    ts[0].text = new_caption
    for t in ts[1:-1]:
        t.text = ""


def set_tof(doc, cur, new):
    _set_tof_caption(find_tof(doc, cur)._p, new)
    print(f"OK TOF: {cur[:8]} -> {new[:10]}")


def delete_tof(doc, cur):
    p = find_tof(doc, cur)
    p._p.getparent().remove(p._p)
    print(f"OK TOF delete: {cur[:12]}")


def insert_tof_clone_before(doc, ref_cur, new_caption):
    ref = find_tof(doc, ref_cur)
    clone = copy.deepcopy(ref._p)
    ref._p.addprevious(clone)
    _set_tof_caption(clone, new_caption)
    print(f"OK TOF insert: {new_caption[:10]}")


# --------------------------------------------------------------------------- #
# new-table content (all numbers traceable, see module docstring)
# --------------------------------------------------------------------------- #
DEPLOY_CAP = "圖二、工程部署架構圖"
DEPLOY_DESC = (
    "圖二補充圖一之概念分層，呈現可實際部署之工程路徑。GitHub Pull Request 由 GitHub Actions "
    "或命令列介面觸發，經 TLS 反向代理呼叫 FastAPI 非同步工作端點，再由多階段 CoT 管線協調 "
    "Gemma-4-31B-it 推論。LoRA 適配器可於啟動時選擇掛載或停用，RAG 則由 "
    "EmbeddingGemma 建立 FAISS 索引，並從具雜湊紀錄之版本化規則語料取回規則。各階段輸出與"
    "逐行 findings 回存為可追溯產物，核心推論服務部署於 DGX Spark／GB10 目標主機。"
)

DS_CAP = "表二、基準資料集統計"
DS_ROWS = [
    ["項目", "值"],
    ["案例數", "44（資料夾標記 ChatGPT 22、Copilot 22，確切版本未留存）"],
    ["類別分布", "bad_data 4、code_diff 12、only_code 28"],
    ["非空白程式行數", "總計 3241，平均 73.7，中位數 65，最小 37，最大 150"],
    ["標註真問題數", "總計 333，平均每案 7.6，最小 1，最大 17"],
    ["問題嚴重度分布", "critical 7、high 60、medium 181、low 85"],
    ["原始碼唯一性", "only_code 之 ChatGPT 第七與第八案位元組相同，去重後為 43 筆相異案例"],
]
DS_DESC = (
    "表二逐項統計本研究之 44 筆基準資料集。全體共 3241 行非空白程式碼，單案自 37 行至 "
    "150 行不等，平均 73.7 行。為支援客觀問題偵測評估，本研究另以 gpt-5.6-sol 僅讀原始碼、"
    "獨立列出每案應被指出之真實問題作為標註基準，全體合計 333 個真問題，平均每案 7.6 個，"
    "並依 critical、high、medium、low 標註嚴重度。須誠實指出者，此一標註基準由大型語言模型"
    "產生，屬 model-derived reference 而非人工黃金標準，其人工複核列為未來工作（見 §6.3）。"
    "另需揭露，only_code 類別中由 ChatGPT 生成之第七與第八案原始碼位元組完全相同，屬重複"
    "案例，44 筆去重後實為 43 筆相異原始碼，逐案評測時仍按 44 筆計算，以與各條件之逐案"
    "輸出對齊。"
)

ABSTRACT_METHOD_ZH = (
    "將 LLM 直接用於審查時，會遇到三個彼此牽連之問題：模型可能產生幻覺（編造看似合理卻"
    "不正確之內容）、相同輸入下輸出不穩定而一致性不足、以及缺乏對特定專案領域規範之認識。"
    "本研究因此提出一套整合框架：以多階段思維鏈拆解審查，以檢索增強生成動態注入專案規則，"
    "再以模型生成之回答資料進行監督式微調。現行 Gemma-4-31B-it 訓練入口預設採 bf16 LoRA，前代 Qwen "
    "配置另採 QLoRA。作者確認教師模型為 ChatGPT 5.4，惟其參數規模未公開，且訓練資料未保存"
    "逐筆生成請求與回應之 provenance，因此本研究將此訓練定位為 "
    "response-based distillation／SFT 管線，不宣稱已證明由特定大型教師向較小學生模型作尺寸壓縮。"
    )

ABSTRACT_METHOD_EN = (
    "Applying LLMs directly to code review raises three intertwined problems: hallucination, unstable "
    "and inconsistent outputs, and limited awareness of project-specific norms. The framework combines "
    "a multi-stage review workflow, retrieval-augmented rule injection, and supervised fine-tuning on "
    "model-generated answers. The current Gemma-4-31B-it training entry point defaults to bf16 LoRA, whereas the "
    "previous Qwen configuration used QLoRA. The author identifies the teacher as ChatGPT 5.4; however, "
    "its parameter scale is not public and per-record generation provenance was not retained. This study "
    "therefore describes the training component as a response-"
    "based distillation/SFT pipeline and does not claim verified size compression from a specific larger "
    "teacher into a smaller student."
    )

PROBLEM_MAPPING = (
    "對應地，本研究以三項設計分別處理流程完整性、領域規範與部署成本。多階段工作流將審查拆解"
    "為摘要、初步審查、靜態分析、程式碼異味偵測與最終彙整，RAG 於生成前注入命中之專案規則，"
    "參數高效微調則以模型生成回答作 SFT。現行 Gemma 訓練入口預設採 bf16 LoRA，前代 Qwen 才採 QLoRA。"
    "教師模型由作者確認為 ChatGPT 5.4，但未保存逐筆生成 provenance，且模型參數規模未公開。"
    "因此第三項貢獻限於可驗證之資料管線與 LoRA 訓練工程，不推論特定"
    "教師之能力或參數規模已被壓縮移轉。"
    )

KD_SIGNIFICANCE = (
    "對本研究的意義：現存資料與程式碼可證明，本研究以模型生成之 question／answer 資料對學生"
    "模型進行 SFT，且只將 answer token 納入 loss，think 欄未使用。此作法可稱為 response-based"
    " distillation。作者確認教師模型為 ChatGPT 5.4，但資料未保存逐筆生成 provenance，且教師"
    "參數規模未公開，故不能證明由特定較大型教師向較小學生模型作尺寸壓縮，也不宣稱學生權重"
    "習得 think 欄之推理軌跡。多階段審查之"
    "可追溯性來自推論管線逐步落檔，而非權重內化之推理。"
    )

DATASET_SOURCE = (
    "本研究使用 44 筆 Python 合成測試案例。程式庫目錄標記其中 22 筆為 ChatGPT、22 筆為"
    " Copilot，惟未保存兩項工具之確切模型版本、生成參數與完整提示 provenance，也未保留可核對"
    "之人工驗證程序。故來源標籤只用於描述資料夾分組，不作為資料品質或代表性已獲證明之依據。"
    )

DATASET_SIZE = (
    "44 筆案例涵蓋 bad_data 4 筆、code_diff 12 筆與 only_code 28 筆，完整統計列於表二。此數量"
    "為現存實驗資料之規模，不主張 44 筆為最佳樣本數，也不據此推論已充分涵蓋所有錯誤類型。"
    )

DATASET_SCOPE = (
    "這 44 筆資料用於在相同輸入下比較多階段與單一提示詞。資料皆為合成 Python 案例，且其中"
    "兩案原始碼重複，去重後為 43 筆相異案例，因此結果僅適用於此基準，不能直接外推至真實 PR"
    " 或其他程式語言。"
    )

OBJ_CAP = "表八、客觀問題偵測評估（多階段流程與單一提示詞）"
OBJ_ROWS = [
    ["指標", "多階段流程", "單一提示詞"],
    ["Recall（召回率）", "0.592（197／333）", "0.498（166／333）"],
    ["Severity-weighted recall（嚴重度加權召回率）", "0.602", "0.510"],
    ["Precision（精確率）", "0.988（474／480）", "0.988（329／333）"],
    ["F1", "0.740", "0.663"],
    ["平均每案宣稱問題數", "10.9", "7.6"],
    ["correctness 類召回率", "0.489（67／137）", "0.387（53／137）"],
    ["bug 類召回率", "0.522（35／67）", "0.403（27／67）"],
    ["security 類召回率", "0.571（12／21）", "0.476（10／21）"],
    ["smell 類召回率", "0.815（22／27）", "0.630（17／27）"],
    ["maintainability 類召回率", "0.763（61／80）", "0.738（59／80）"],
    ["design 類召回率", "0.000（0／1）", "0.000（0／1）"],
    ["逐案召回率比較（勝／平／負）", "26／14／4", "4／14／26"],
]
OBJ_DESC = (
    "本表於 §4.1 之 44 筆基準（合計 333 個標註真問題）上，以客觀問題偵測指標比較多階段"
    "流程與單一提示詞。多階段流程之召回率明顯較高（0.592 對 0.498），二者精確率相同"
    "（皆 0.988），故 F1 亦較高（0.740 對 0.663）。表中並列六類問題之召回率與逐案比較，"
    "多階段於 correctness、bug、security、smell 與 maintainability 五類均較高，design 類唯一"
    "一題則兩者皆未命中。逐案比較 44 筆，多階段召回較高者 26 案、相同者 14 案、較低者 4 案。"
)

JUDGE_CAP = "表九、LLM-as-a-Judge 五維主觀評分之多階段與單一提示詞差異"
JUDGE_ROWS = [
    ["維度", "Claude 裁判 Δ", "gpt-5.6-sol 裁判 Δ"],
    ["Readability", "−2.3", "−0.3"],
    ["Maintainability", "−6.9", "−0.5"],
    ["Correctness", "−0.3", "+1.5"],
    ["Multi-Review Coverage", "−9.3", "−7.7"],
    ["Comprehensiveness", "−6.3", "−1.0"],
    ["五維平均 Δ", "−5.01", "−1.60"],
]
JUDGE_DESC = (
    "本表以兩位獨立裁判（Claude 與 gpt-5.6-sol）於五維百分制下評分，欄位 Δ 為多階段流程"
    "減單一提示詞之分差（n=44），負值代表單一提示詞得分較高。兩位裁判之整體均分為："
    "Claude 下多階段 78.8、單一提示詞 83.8，gpt-5.6-sol 下多階段 84.4、單一提示詞 86.0。"
    "可見於主觀行文品質上單一提示詞相當或略佳，唯 Correctness 一維多階段打平或略勝，"
    "此與表八之客觀偵測結論互補而非矛盾。"
)

HUMAN10_CAP = "表十、前代 Qwen 配置之人工評分：多階段提示詞與單一提示詞比較"
HUMAN11_CAP = "表十一、前代 Qwen 配置之人工評分：微調與未微調比較"

TRAIN_ROWS = [
    ["參數", "配置"],
    ["設定證據層級", "以下為現行入口之程式預設，未保存該次執行環境快照"],
    ["基座模型", "評估報告標示 google/gemma-4-31B-it，現行入口亦以此為預設"],
    ["微調方法", "現行入口預設 bf16 LoRA（非 4-bit QLoRA）"],
    ["實際訓練資料版本", "無法確定，未保存資料 hash、快照或 console log"],
    ["程式庫資料歷史", "修復前 663 個實體行且 8 行串接多物件，現行 695 筆有效 JSON，其中 5 筆完全重複"],
    ["資料切分", "載入成功之資料全作 train，validation 0、test 0，實際 train 筆數無法確定"],
    ["訓練平台證據", "Compose／Dockerfile 目標為 DGX Spark GB10，未保存該次硬體快照"],
    ["SEQ_LEN", "1024"],
    ["MICRO_BATCH_SIZE", "1"],
    ["GRADIENT_ACCUMULATION_STEPS", "64（有效 batch size 64）"],
    ["NUM_EPOCHS", "3"],
    ["LEARNING_RATE", "2e-5"],
    ["WARMUP_STEPS", "依程式預設計算為 3，實際 run 未留 log"],
    ["WEIGHT_DECAY", "0.01"],
    ["LORA_R／ALPHA／DROPOUT", "64／64／0.1"],
    ["optimizer／scheduler", "adamw_torch_fused／cosine"],
    ["optimizer steps", "依兩個候選資料筆數與預設參數皆計算為 33，實際 run 未留 log"],
    ["資料 shuffle seed", "42"],
    ["提示模板隨機性", "3 種模板以 random.choice 選取，未設定 Python random seed"],
    ["監督訊號", "僅 answer token 計算 loss，think 欄未使用"],
    ["logging", "TensorBoard 每 1 optimizer step，現存庫未保留事件檔"],
    ["checkpoint policy", "每 20 steps、最多 3 份、另存 final，無 best-model 選擇"],
]

TRAIN_EXPLAIN_ROWS = [
    ["參數", "意義說明"],
    ["訓練版本／程式庫資料歷史", "評估報告時間早於修復提交不能排除未提交變更，故不據此反推實際訓練筆數"],
    ["資料切分", "訓練腳本只建立 train dataset，未建立 validation 或 test dataset"],
    ["SEQ_LEN", "每筆輸入含 BOS／EOS 後的最大序列長度"],
    ["MICRO_BATCH_SIZE", "單次前向與反向傳遞的樣本數"],
    ["GRADIENT_ACCUMULATION_STEPS", "累積 64 次 micro batch 後更新一次參數"],
    ["NUM_EPOCHS", "完整走訪訓練資料的次數"],
    ["WARMUP_STEPS", "33 個 optimizer steps 中前 3 步作學習率預熱"],
    ["LORA_R／ALPHA／DROPOUT", "低秩維度、縮放係數與 LoRA dropout"],
    ["提示模板隨機性", "模板選擇未固定 Python random seed，實際逐筆模板與 token 總數不能精確回推"],
    ["監督訊號", "prompt 標為 −100，只有 answer token 與 EOS 納入 loss"],
    ["checkpoint policy", "未提供 eval dataset，故沒有依驗證 loss 選擇最佳 checkpoint"],
]

TRAIN_GENERATION_ROWS = [
    ["參數名", "設定"],
    ["現行 Gemma 訓練入口預設", "bf16 LoRA"],
    ["訓練權重量化", "無（未使用 BitsAndBytes／NF4）"],
    ["LoRA target", "language_model 之 q／k／v／o／gate／up／down projection"],
    ["gradient checkpointing", "True（use_reentrant=False）"],
    ["adapter 輸出", "outputs-lora-gemma4-31b，未合併"],
    ["現行 Gemma 服務權重精度", "Compose 預設 FP8，可覆寫 bf16，實驗未留存啟動環境快照"],
    ["注意力實作", "SDPA（server image 移除 flash-attn）"],
    ["do_sample", "False（貪婪解碼）"],
    ["temperature／top_p", "未使用"],
    ["max_new_tokens", "8192（本輪實驗請求值）"],
    ["repetition_penalty", "未設定"],
    ["seed", "不適用於貪婪解碼"],
    ["前代 Qwen 訓練", "另有 4-bit NF4 QLoRA 腳本，不代表現行 Gemma 設定"],
]

PACKAGE_ROWS = [
    ["套件／環境", "版本需求", "用途"],
    ["NGC PyTorch image", "nvcr.io/nvidia/pytorch:25.09-py3", "DGX Spark GB10 訓練基底"],
    ["transformers", ">=5.10,<6", "載入 Gemma 4、Trainer 與生成"],
    ["peft", ">=0.17", "LoRA 適配器"],
    ["datasets", ">=3.0", "載入 JSONL 訓練資料"],
    ["accelerate", ">=1.0", "訓練裝置與執行支援"],
    ["sentence-transformers", ">=3.0", "EmbeddingGemma RAG 嵌入"],
    ["tensorboard", "未鎖定", "訓練 loss 記錄介面"],
    ["bitsandbytes", "現行 Gemma 未安裝／未使用", "僅前代 Qwen QLoRA 腳本使用"],
]


# --------------------------------------------------------------------------- #
# prose rewrites (full paragraph)  (anchor -> new text)
# --------------------------------------------------------------------------- #
RQ2 = (
    "RQ2：在固定參數規模與相同基座下，僅以多階段思維鏈提示詞拆解審查任務、而不進行模型"
    "微調，能否較單一提示詞偵測出更多真實問題，於問題偵測之召回率與 F1 上取得提升，藉以"
    "分離提示詞流程與權重微調二者之獨立貢獻？"
)
RQ3 = (
    "RQ3：多階段流程相對單一提示詞之邊際貢獻主要落在何處，是客觀之問題偵測完整性"
    "（召回率與 F1），抑或主觀之審查意見行文品質？"
)
MOTIVATION_REFRAME = (
    "這三個問題彼此關聯，因而需要從流程、領域規範與部署三個層次分別處理。將全部審查任務置於"
    "單一提示詞時，模型須在一次推論中兼顧摘要、風格、邏輯與安全，可能分散注意力並遺漏問題，"
    "缺乏專案規範亦可能使審查停留於通用建議。據此，本研究以多階段流程處理問題偵測，以 RAG "
    "設計提供專案規則接地，以知識蒸餾結合 LoRA 支援有限資源部署，前代 Qwen 配置另使用 QLoRA。"
    "三項設計之角色不同，其獨立"
    "效益是否成立須分別驗證，不能僅由整合架構之存在推定。"
)
OBJECTIVE_ABLATION = (
    "為釐清整合框架內部各元件之相對作用，本研究比較多階段流程與單一提示詞，並嘗試分離模型"
    "微調與 RAG 之影響。現行資料可直接回答多階段流程之問題偵測完整性，RAG 因 44 案皆未命中"
    "檢索閾值而形成 null 結果，Gemma 微調則因缺少可切換之基礎模型端點而無法乾淨消融。第五章"
    "據此區分已實證結果與尚待驗證之元件效益。"
)
CONTRIBUTION_INTRO = (
    "本研究之貢獻包含多階段結構化審查工作流、RAG 領域規則注入與知識蒸餾結合 LoRA 之部署"
    "設計。三者分別對應問題偵測、專案規範接地與有限資源部署，但目前量化證據強度並不相同："
    "多階段流程已有本論文之客觀問題偵測評測，RAG 於本合成基準未產生命中，現行 Gemma 微調之"
    "乾淨消融亦未完成。下列七項貢獻因此分為已實證之流程結果與尚待後續量化之設計／工程成果。"
)
COT_ROLE = (
    "在整條工作流中的角色：本步驟承接前序發現並產生可供後續彙整之文字理由，使每一階段產物可"
    "獨立檢視與追溯。這種可追溯性來自流程結構與落檔產物，不表示文字理由忠實映射模型內部推理，"
    "亦不能單獨作為幻覺率下降或忠實可解釋性提升之證據。"
)
COT_CAPTION = (
    "本圖為套用 Chain-of-Thought 提示詞後之階段性說明範例，用以呈現可檢視之中間產物，不將其"
    "視為模型真實內部推理或忠實可解釋性之直接證據。"
)
CONTRIBUTION_CLOSE = (
    "此外，本研究隨附之開源框架已將上述審查流程整合並部署於 GitHub Actions CI/CD，另實作 "
    "JudgeStep 裁決至 GitHub Review event 之映射、四類推論後端、MCP 整合層與十七項研究級"
    "擴充機制。這些工程成果提供後續實證之基礎，其對真實團隊之審查時間、誤報率、採納率與軟體"
    "品質之效益尚未量測，故不在本論文中推論已能提升實際開發效率或品質。"
)

S1_INTRO = (
    "為釐清本研究所提出之程式碼審查框架中，各關鍵設計對生成結果之影響，本節彙整四類評估"
    "結果：表七之 CRSCORE++ 三維度用於與既有基準作描述性對照，表八之客觀問題偵測用於量測"
    "多階段流程與單一提示詞實際找出多少真問題，表九之 LLM-as-a-Judge 五維主觀評分用於並陳"
    "兩者之審查意見行文品質，表十與表十一之前代 Qwen 人工評分則用於檢查前代配置下自動"
    "評估之趨勢能否被人類評審重現。"
)

LIT_JUDGE = (
    "對本研究的意義：LLM-as-a-Judge（讓強大之 LLM 依指定標準替其他模型輸出打分之方法）使"
    "本研究得以在不依賴大量人工標註之前提下，比較多階段與單一提示詞之審查品質。現行 Gemma "
    "重放採 Claude Opus 4.8 與 gpt-5.6-sol 雙裁判，並將其主觀評分與客觀問題偵測結果並陳，"
    "前代 Qwen 配置則保留 GPT-5、Gemini-3 與人工評分作歷史交叉驗證。此分層安排用以揭露裁判"
    "模型與輸出表示法之敏感性，而不把單一 LLM 裁判視為人工黃金標準。"
)

S522_INTRO = (
    "表九聚焦於現行 Gemma-4-31B-it 重放之主觀審查品質。受評條件為同一基座、同一份 44 筆"
    "資料與相同貪婪解碼設定下的多階段流程與單一提示詞，裁判為 Claude Opus 4.8 與 "
    "gpt-5.6-sol。兩位裁判皆依相同五維規準評讀最終審查文件，表內以多階段減單一提示詞之分差"
    "呈現，藉此直接觀察裁判與輸出表示法對主觀品質判定之影響。"
)
S522_DESIGN = (
    "兩個受評條件只改變審查流程。多階段條件依序產生初步摘要、初步審查、Linter、程式碼異味"
    "與總彙整，評分時使用最終總彙整文件，單一提示詞條件則於一次推論中完成審查並直接評分其"
    "單一輸出。此比較量測的是最終文件之可讀性、可維護性、正確性、跨評論覆蓋與完整性，不等同"
    "於表八以所有階段 findings 計算的客觀問題偵測召回率。"
)
S522_RESULT = (
    "如表九所示，兩位裁判均使單一提示詞之五維平均略高於多階段流程，Claude 之平均分差為"
    "負 5.01 分，gpt-5.6-sol 為負 1.60 分，且各維分差幅度並不一致。這表示主觀行文評分對"
    "裁判模型與輸出形式敏感，單一提示詞因形成單一連貫文件並常附重構範例而在呈現上較占優勢，"
    "多階段之主要實證優勢則在表八所示之問題偵測完整性。"
)
S51_RELATION = (
    "表八與表九回答不同問題。表八以標註真問題量測偵測成效，顯示多階段流程在召回率與 F1 較高"
    "且精確率不降，表九則由兩位裁判評讀最終文件之主觀品質，顯示單一提示詞在呈現上相當或略佳。"
    "兩表並陳可避免把找到更多真問題與文字讀來較順誤當成同一種品質。"
)

CRSCORE_DESC = (
    "本表以 CRSCORE++ 評分法在 comprehensiveness、conciseness、relevance 三個面向上，列出"
    "本研究三個學生模型與 CRSCORE++ 基準之整體表現。表中三個 Ours 欄均由 Claude 以標準 "
    "CRSCORE++ 評分提示詞、於本研究之基準資料上評分，CRSCORE++ 基準欄則沿用其原始論文之 "
    "GPT-4o-mini 分數，二者之裁判模型與受評資料皆不同，故表七僅適合作為描述性對照，跨欄之"
    "高下比較僅屬指示性參考，可讀性、正確性與可維護性等細部差異則交由表八之客觀偵測與"
    "表九之五維主觀評分診斷。"
)

HUMAN10_DESC = (
    "本表為前代 Qwen 學生模型配置（2026-02 之 Qwen3-Coder-30B-A3B）之人工評分結果，比較"
    "多階段提示詞與單一提示詞之審查品質，用以於前代配置上交叉驗證自動評分與人類判斷之一致"
    "性。CSV 共含 8 份有效回覆，多階段微調組之五維平均為 85.6，單一提示詞組為 79.25，逐份"
    "回覆比較時前者為 6 勝、2 負。本表屬前代配置之人工評估，不與本輪 Gemma 之客觀偵測結果"
    "（表八）混用。"
)
HUMAN11_DESC = (
    "本表為前代 Qwen 學生模型配置之人工評分結果，比較微調（Finetune）與未微調"
    "（Without Finetune）學生模型之審查品質。本表屬前代配置之人工評估，其所示之微調小幅"
    "改善僅適用於前代 Qwen 配置。兩個多階段組之五維平均為 85.6 與 85.4，逐份回覆比較各為"
    " 4 勝、4 負，不能據此宣稱微調穩定改善整體品質。現行 Gemma 之 base-model 部署開關已完成"
    "實作，但 44 案重放尚未完成，詳見 §6.3。"
)

HUMAN_PROTOCOL = (
    "雖已採用 LLM-as-a-Judge 進行大規模自動評估，本研究仍以前代 Qwen 配置之人工問卷作歷史"
    "補充。作者確認現存 CSV 之 8 份有效回覆分別來自 8 名不同評分者，每人均對三組審查輸出"
    "各評五個維度。CSV 未含評分者識別碼，故人數係以作者確認為依據，不能由檔案獨立核驗身分。"
    "評分手冊要求評分者自行"
    "隨機挑選一個結果資料夾，再比較同名之三組輸出，CSV 卻未記錄案例編號，因而無法確認每份"
    "回覆所評案例，也無法計算正式之評分者間信度。問卷與手冊皆未記錄專業背景或開發年資，"
    "方法資料夾名稱又揭露微調與流程條件，故本研究不宣稱採盲評。可確認的隨機程序僅有案例"
    "自行抽選，三組方法之呈現順序未見隨機化紀錄。"
)

HUMAN_CALC = (
    "分數彙整時，先對同一受評方法、同一維度下之 8 份回覆取平均，再將各維度平均填入表十與"
    "表十一。除平均分數外，本研究另在每份回覆內作成對方向比較，以避免不同評分者可能抽到"
    "不同案例所造成之誤解。微調多階段組相對單一提示詞組之五維平均為 6 勝、2 負，其中建設性"
    "為 7 勝、1 負，正確性為 5 勝、1 平、2 負，覆蓋與可抽取性為 5 勝、2 平、1 負。相較之下，"
    "微調與未微調多階段組之五維平均各為 4 勝、4 負，平均僅差 0.2 分。由於 CSV 未記錄案例"
    "編號與評分者識別碼，本研究不計算 Cohen's κ 或 ICC，亦不把這 8 份回覆描述為正式盲評。"
)

S42_INTRO = (
    "本研究的評估分為多個層次，目的不是重複計分，而是分別回答不同問題。CRSCORE++ 三維度"
    "沿用既有研究之完整性、簡潔性與相關性指標，使本研究可與基準方法作描述性對照。本研究"
    "設計之五維度百分制則細看審查意見是否清楚、可執行、正確且能覆蓋多份分析結果。除此二層"
    "主觀評分外，本研究另加入一層客觀之問題偵測評估，以標註之真實問題為基準，量測各流程"
    "實際找出多少真問題。最後以前代 Qwen 配置之人工評分交叉驗證，確認趨勢不是單一裁判模型"
    "或提示詞造成的偏差。換言之，表七以 CRSCORE++ 三維度與既有基準作描述性對照，表八以"
    "客觀問題偵測回答「多階段流程與單一提示詞何者找出更多真問題」，表九以五維百分制之 "
    "LLM-as-a-Judge 並陳兩者之主觀行文品質，表十與表十一則為前代 Qwen 配置之人工評分，供"
    "人類判斷與自動評分之交叉參照。"
)
S42_OBJ_METHOD = (
    "客觀問題偵測評估之流程如下：以 §4.1 所述之標註真問題為基準，多階段流程之宣稱問題取自"
    "其五步驟之完整 findings（初步摘要、初步審查、Linter、程式碼異味與總彙整），單一提示詞"
    "則取自其單次審查輸出。對每一筆審查結果，由 Claude 逐條比對其所提意見是否命中標註之"
    "真問題，命中數除以該案真問題總數即為召回率，全體 44 案加總後之召回率反映流程找出多少"
    "真正該提的問題。另以嚴重度加權召回率依 critical、high、medium、low 給予遞減權重，衡量"
    "增益是否涵蓋較嚴重之問題。精確率則為流程所宣稱之問題中經判定成立者之比例，反映誤報"
    "多寡，F1 為召回率與精確率之調和平均。此組客觀指標與前述主觀評分互補：主觀評分衡量"
    "審查意見讀來如何，客觀偵測衡量究竟找出多少真問題。"
)

S531_A = (
    "本節將上述各表串成一條論證，並逐一回指 §1.2 之三大問題與 §1.4 之三項設計。先看自動"
    "評分與人工評分之一致性，對應 RQ4。由於本輪 Gemma 重放未再執行人工評分，此處之人類"
    "交叉驗證以前代 Qwen 學生模型配置之人工評分（表十、表十一）為據：於該配置下，套用多"
    "階段提示詞與微調後，Maintainability、Correctness 與 Multi-Review Coverage 三維度之人工"
    "分數均高於基礎模型。微調與未微調之比較則僅 Maintainability 與 Correctness 略高，"
    "Readability、Multi-Review Coverage 與 Comprehensiveness 均未呈一致改善。因此前代人工結果"
    "只能支持多階段流程於部分深層維度較佳，以及微調於建設性與正確性有小幅正向差異，不能概括"
    "為微調提升整體品質。此差異亦與 [11] 所述 "
    "LLM-as-a-Judge 於語感類指標相對人類較寬鬆之偏誤一致，指出一個明確之後續校正方向。"
)
S531_B = (
    "由上述前代交叉驗證可知，於 Qwen 世代，兩套獨立評審在多數深層維度上指向同一結論，"
    "提供有限度之歷史交叉驗證。須誠實界定者，"
    "此一致性係於前代 Qwen 配置上取得，本輪更換基底為 Gemma-4-31B-it 後之人工複評尚未執行，"
    "列為研究限制與未來工作，詳見 §6.3。"
)

S532_A = (
    "次看多階段流程相對單一提示詞之邊際貢獻，對應 RQ2 與 RQ3。就客觀之問題偵測而言，如表八"
    "所示，於 44 筆基準之 333 個標註真問題上，多階段流程找出 197 個（召回率 0.592），單一"
    "提示詞找出 166 個（召回率 0.498），多階段多找出 31 個真問題，相對多出約百分之十九，"
    "召回率高出 9.3 個百分點。二者之精確率相同，皆為 0.988，代表其所宣稱之問題幾乎都成立、"
    "誤報極少，故綜合指標 F1 亦較高（0.740 對 0.663）。以嚴重度加權後之召回率呈相同方向"
    "（0.602 對 0.510），顯示多階段之增益不僅在數量，亦涵蓋較嚴重之問題。"
)
S532_B = (
    "如表八所示，逐類觀之，多階段流程於五個有命中之類別召回率皆高於單一提示詞：correctness "
    "為 0.49 對 0.39、bug 為 0.52 對 0.40、security 為 0.57 對 0.48、smell 為 0.81 對 "
    "0.63、maintainability 為 0.76 對 0.74，design 類唯一一題則兩者皆未命中。逐案比較 44 筆，"
    "多階段召回較高者 26 案、與單一"
    "提示詞相同者 14 案、較低者僅 4 案。此一優勢之來源在於多階段流程設有專屬之 Linter 與 "
    "Code Smell 階段，能系統性覆蓋單一提示詞於一次推理中易遺漏之規則類與結構類問題，多階段"
    "平均每案提出 10.9 條問題、單一提示詞為 7.6 條，且其多提出者多為真問題，故精確率未因"
    "宣稱數增加而下降。就 RQ2 與 RQ3 而言，多階段流程之邊際貢獻主要體現於問題偵測之完整性，"
    "亦即以相同精確率找出更多真問題，而非審查意見之行文修飾。"
)
S532_C = (
    "須並陳的是，若改以 LLM-as-a-Judge 評估審查意見之主觀行文品質，結論並不相同。如表九"
    "所示，於五維百分制下，兩位獨立裁判之評分反而略偏向單一提示詞：Claude 裁判之五維平均"
    "差為每維約負 5.01 分（多階段整體均 78.8、單一提示詞 83.8），gpt-5.6-sol 裁判為每維約"
    "負 1.60 分（84.4 對 86.0），差距主要來自 Multi-Review Coverage 與 Comprehensiveness "
    "兩維，唯 Correctness 一維多階段打平或略勝。此一並陳並不矛盾，而是反映兩種評估回答不同"
    "問題：客觀偵測衡量找出多少真問題，多階段因專項分工而勝出，主觀評分衡量審查意見讀來"
    "如何，單一提示詞因輸出為單一連貫文件、且常直接附上重構範例而於呈現略占優勢。又因 "
    "LLM-as-a-Judge 之評分對輸出表示法與裁判模型敏感，兩位裁判之分差幅度即不相同，宜與客觀"
    "指標並看而非擇一。至於模型微調之邊際貢獻，本輪已完成之客觀評測仍使用掛載 LoRA 的 "
    "Gemma-4-31B-it，故未能乾淨消融。服務程式現已新增可停用 LoRA 之 base-model 部署開關，"
    "但相同 44 案之重放尚未完成，不能先行宣稱微調之獨立效益。前代 Qwen 人工評分（表十、"
    "表十一）僅作歷史參考。"
)

S533_A = (
    "末看與基準方法之描述性對照，對應 RQ1。如表七所示，以同一裁判（Claude）、標準 "
    "CRSCORE++ 評分提示詞評分時，本框架以 Gemma-4-31B-it 與前代學生模型 Qwen3-Coder-30B-"
    "A3B 為基底，二者於三維度相當（comprehensiveness 皆 1.00，conciseness 為 0.79 對 "
    "0.78，relevance 皆 0.86），顯示更換基座未損及既有審查品質。三個 Ours 欄於三維度亦高於 "
    "CRSCORE++ 基準之原始報告分數（0.67／0.57／0.63），小型變體 Qwen2.5-Coder-7B 為 "
    "0.97／0.69／0.79，完整性仍近飽和，簡潔性與相關性則明顯低於前代基底。須說明的是，基準"
    "欄沿用 CRSCORE++ 原始論文之 GPT-4o-mini 分數，其裁判模型與受評資料皆與 Ours 欄不同，"
    "故表七僅作描述性對照，跨欄比較僅屬指示性，RQ1 不以表七逕行作答，多階段流程之實質效益"
    "改以表八之客觀問題偵測為據。"
)
S533_B = (
    "綜合 §5.3.1 至 §5.3.3，本研究之客觀評測支持一項明確結論：於 44 筆基準之問題偵測任務上，"
    "多階段流程較單一提示詞找出更多真實問題（表八：召回率 0.592 對 0.498、F1 0.740 對 "
    "0.663），且精確率不降。至於主觀行文品質，兩者相當而單一提示詞於呈現略佳，二項指標"
    "互補，多階段贏在問題偵測之完整性，單一提示詞贏在呈現。RAG 檢索稽核顯示，前次採前代"
    "嵌入閾值 0.7 時為零命中，改用 EmbeddingGemma 校準閾值 0.32 後則有 25／44 案命中，合計"
    "取回 133 份規則文件。此結果證明檢索不再為空，但生成品質之獨立效益仍須完成同條件重放"
    "後判定。模型微調之 base-model 部署開關已完成，但相同 44 案尚未重放，故本"
    "研究不將三項元件宣稱為皆屬必要條件，而是分別定位其角色，多階段流程提供問題偵測之完整性，RAG 於具"
    "專案語料時提供領域規範之接地，response-based distillation 結合 LoRA 提供有限資源下之"
    "參數高效訓練能力，三者於"
    "整合框架中互補。"
)

C61_INTRO = (
    "本研究提出一套整合多階段結構化審查、RAG 規則檢索與參數高效微調之程式碼審查框架。"
    "三項元件之證據強度不同：多階段流程已於 44 案基準量得較高之問題偵測召回率與 F1，RAG "
    "已確認於校準閾值下可取回規則，但其生成品質影響仍待完整重放評分，微調則可由程式庫確認"
    "入口、LoRA 管線與 base-model 部署開關，惟現行 Gemma 之獨立消融尚未完成，實際訓練資料"
    "版本亦無法確認。因此，本研究"
    "將框架視為可供分項驗證之整合實作，不把架構存在本身當作幻覺率下降、輸出穩定性提升或"
    "特定教師能力已移轉之證據。"
)

C61 = (
    "經 CRSCORE++ 之描述性對照、客觀問題偵測評測與 LLM-as-a-Judge 五維主觀評分，本框架以 "
    "Gemma-4-31B-it 與前代學生模型 Qwen3-Coder-30B-A3B 為基底，於 CRSCORE++ 三維度相當且"
    "均高於基準之原始報告分數（表七：1.00／0.79／0.86 對 1.00／0.78／0.86），但因基準欄"
    "之裁判與資料不同，此處僅為描述性觀察，不以之回答 RQ1。就客觀問題"
    "偵測而言，於 44 筆基準之 333 個標註真問題上，多階段流程較單一提示詞找出更多真問題"
    "（表八：召回率 0.592 對 0.498、F1 0.740 對 0.663）且精確率相同（皆 0.988），此為多"
    "階段流程於問題偵測完整性上之邊際貢獻。至於主觀行文品質，兩位獨立裁判之五維評分顯示"
    "單一提示詞相當或略佳（表九），故本研究誠實並陳二項互補之結論：多階段流程贏在找出更多"
    "真問題，單一提示詞贏在意見之呈現。本研究之核心主張因此並非以單一元件取代其餘，而是"
    "多階段流程、RAG 規則檢索與 response-based distillation 結合 LoRA 於整合框架中各司其職。"
    "其中只有多階段之問題偵測完整性已由本輪客觀比較支持，RAG 生成效益與微調獨立效益仍須"
    "另行完成消融，不由整合架構本身推定。"
)

C62_1 = (
    "本研究之核心貢獻（已於 §5 達成並具量化證據，與 §1.5 所列「將提出」之設計貢獻相區分）"
    "可歸納為三項，三者構成一套整合框架而非彼此可替代之選項。其一，於審查流程設計上，提出"
    "將審查拆解為摘要、初步審查、靜態分析與程式碼異味偵測之多階段思維鏈流程，並以 "
    "build_global_rule_template 統一注入七條靜態全域規則與 RAG 檢索所得之條件式第八條領域"
    "規則。於 44 筆基準之客觀問題偵測評測中，此一流程設計相對單一提示詞以相同精確率找出"
    "更多真實問題（表八：召回率 0.592 對 0.498、F1 0.740 對 0.663），為其於問題偵測完整性"
    "上之邊際貢獻。RAG 與 response-based distillation 結合 LoRA 為框架中另兩項可分離設計，"
    "其獨立品質效益不由多階段結果推定，分別受本輪 RAG 重放與微調消融限制約束。"
)
C62_3 = (
    "其三，於評估方法上，提出 LLM-as-a-Judge-Our 五項百分制細粒度指標，並以多個獨立評審"
    "交叉檢查評分敏感性。現行 Gemma 重放以 Claude 與 gpt-5.6-sol 雙裁判並陳，前代 Qwen 配置"
    "之 GPT-5、Gemini-3 與人工評分僅作歷史交叉驗證，不與現行客觀結果混為同一批證據。"
)

LIM2 = (
    "（2） 評審模型偏差：現行 Gemma 重放之主觀評分由 Claude Opus 4.8 與 gpt-5.6-sol 執行，"
    "客觀問題覆蓋亦由 Claude 比對 model-derived reference，仍可能承載裁判模型自身之問題偏好。"
    "前代 Qwen 配置雖另有 GPT-5、Gemini-3 與人工評分可作歷史交叉驗證，但資料、模型與評分批次"
    "不同，不能視為現行結果之人工複評。以同一批 44 案建立人工黃金標準並盲評，列為未來工作。"
)
LIM7 = (
    "（7） 人工評分紀錄之完整性：作者確認前代 Qwen 人工評分之 8 份有效回覆分別來自 8 名不同"
    "評分者，每人對三組輸出各評五個維度。CSV 未保存評分者識別碼、專業背景、開發年資與所抽"
    "案例編號，因此身分不能由檔案獨立核驗，也不能計算正式評分者間信度。評分手冊只要求自行"
    "隨機挑選案例，未"
    "記錄三組呈現順序隨機化，且資料夾名稱揭露實驗條件，故不宣稱盲評。表十與表十一僅作描述性"
    "歷史對照，不能視為現行 Gemma 結果之人工驗證。"
)
LIM8 = (
    "（8） 訓練重現資訊：Gemma 轉接器評估報告標記為 2026-06-11 01：27，JSONL 修復提交則為"
    "同日 12：47。提交前版本有 663 個實體行，其中 8 行串接多個 JSON 物件而非有效 JSONL，"
    "修復後現行檔則為 695 筆有效 JSON，其中 5 筆完全重複。提交時間不能排除修復內容曾以未"
    "提交狀態先供訓練使用，且訓練未保存資料 hash、快照或 console log，故實際訓練究竟載入 "
    "663 或 695 筆均不能由程式庫確定，論文不採任一數字作確定主張。訓練腳本未實作去重、主題"
    "分層或 validation／test 切分，載入成功之資料全作 train。依 663 或 695 兩個候選筆數與現行"
    "入口預設參數計算，皆為 33 個 optimizer steps、3 個 warmup"
    " steps、每 20 steps 儲存且最多保留 3 份，最後另存 final adapter，未設定最佳模型選擇。"
    "教師模型由作者確認為 ChatGPT 5.4，但逐筆生成 provenance 與人工清理紀錄未寫入資料，三種"
    "提示模板又以未固定 seed 之"
    " random.choice 選取，故實際 token 總數無法精確回推。TensorBoard 事件檔、checkpoint 實體、"
    "完整 loss 曲線、峰值 VRAM 與開始／結束 console log 均未保留，實際訓練時間亦無法重建。"
)

LIM1 = (
    "（1） 資料規模與多樣性：程式庫將 44 筆 Python 合成測試資料分為 ChatGPT 與 Copilot 來源"
    "目錄各 22 筆，但未保存確切模型版本、生成參數與完整提示 provenance，故來源名稱僅代表"
    "資料夾標籤。此基準並非真實 PR，其中 only_code 類別之 ChatGPT 第七與第八案原始碼位元組"
    "完全相同，去重後實為 43 筆相異原始碼。相較真實開源專案動輒"
    "上千個 Pull Request 之審查歷史，單一合成基準仍屬有限規模，對其他程式語言（C++、Java、"
    "Go）、其他類型專案（嵌入式系統、分散式系統、前端框架）與真實 PR 之泛化能力尚未驗證。"
)
LIM3 = (
    "（3） 微調範圍與消融：現行 Gemma 訓練入口依表三預設採 bf16 LoRA，前代 Qwen 配置另採"
    " 4-bit QLoRA。現行"
    "本輪已完成之客觀重放使用掛載 LoRA 的 Gemma-4-31B-it，故尚未對微調進行乾淨消融。服務"
    "程式已新增 PRTHINKER_DISABLE_LORA 啟動開關，健康檢查亦會回報 LoRA 狀態，可建立未掛載"
    "適配器之 base-model 組，但相同 44 案尚未完成重放與評分，故微調之獨立效益仍不得推定。"
    "不同基座模型（CodeLlama、StarCoder2 等）於相同 CoT 框架下之系統性比較、"
    "LoRA rank 等超參數消融亦未涵蓋，NF4 對 INT8 僅適用於前代 Qwen 量化路徑。"
)

TRAIN_FLOW = (
    "本研究之現行 Gemma 訓練入口為 docker/docker-compose.train.yml 所呼叫之 codes/train/"
    "gemma4-31b.py。基座模型為 google/gemma-4-31B-it，以 bf16 直接載入並加入未合併之 LoRA"
    " 適配器，非 BitsAndBytes 4-bit QLoRA。程式庫之提交前資料有 663 個實體行，其中 8 行串接"
    "多個 JSON 物件而不是有效 JSONL，修復後現行檔為 695 筆，其中 5 筆完全重複。雖 Gemma "
    "轉接器評估報告之時間早於修復提交，提交時間不能排除修復內容曾以未提交狀態先供訓練使用，"
    "且未保存資料 hash、快照或 console log，故實際訓練筆數無法確定。訓練腳本將載入成功之資料"
    "全部作為 train dataset，未建立 validation 或 test split。每筆包含"
    " Instruction、question、think、answer，訓練時 prompt token 以 −100 遮罩，只有 answer token"
    " 與 EOS 計入 loss，think 欄未進入提示或監督訊號。三種提示模板由 random.choice 選取且未"
    "固定 Python random seed，資料順序則以 seed 42 shuffle，因此實際逐筆模板與 token 總數無法"
    "由現存成果精確回推。現行入口預設 3 epochs、有效 batch size 64、33 個 optimizer steps、3 個"
    " warmup steps、cosine scheduler、adamw_torch_fused 與 gradient checkpointing，但實際 run"
    "之環境覆寫值未保存。Trainer 原設定每一步"
    "寫入 TensorBoard、每 20 步儲存 checkpoint 且最多保留 3 份，最後另存 final adapter，但現存"
    "程式庫未保留事件檔、checkpoint 實體或 console 起訖紀錄。"
    )

QLORA_ROLE = (
    "對本研究的意義：QLoRA 為前代 Qwen 訓練路徑所採用之低記憶體微調方法，相關 4-bit NF4"
    " 與雙重量化設定不應套用至現行 Gemma。現行 Gemma 訓練入口係為 DGX Spark GB10 平台設計，"
    "採 bf16 LoRA 且不使用 bitsandbytes，惟該次執行之硬體快照未保存。故本節保留 QLoRA 作為相關"
    "技術與前代配置說明，現行實驗設定則以表三至表五之 LoRA 資訊為準。"
    )

ARCH_DEPLOY = (
    "線上服務載入 google/gemma-4-31B-it 並掛載未合併之 LoRA 適配器，伺服器映像移除不相容之"
    " flash-attn，使注意力使用 SDPA。Compose 目前以 FP8 為預設權重模式並允許覆寫為 bf16，"
    "但 2026-07 實驗未保存容器啟動環境快照，故論文不宣稱該次服務實際採用何者。離線訓練則"
    "以 bf16 載入同一 Gemma 基座並訓練 LoRA 適配器。"
    )

TRAIN_ENV = (
    "前代 Qwen 實驗之既有環境紀錄為雙 NVIDIA L40S（合計 96 GB VRAM）伺服器。現行 Gemma "
    "訓練與服務路徑則由 docker/docker-compose.train.yml、docker/Dockerfile.train 與 server-"
    "gemma4 組態明示以本機 DGX Spark GB10（aarch64、sm_121）為目標平台，不能將兩代配置混寫"
    "為同一硬體。程式庫未保存該次 Gemma 訓練之 nvidia-smi、容器環境快照、峰值記憶體量測或"
    "主機資產紀錄，因此只能報告目標平台與前代環境，不能獨立確認該次訓練之實際峰值 VRAM。"
)

TRAIN_TABLE_DESC = (
    "本表依 Gemma 現行訓練入口列出程式預設之資料處理、切分、主要超參數、optimizer steps、監督訊號、"
    "logging 與 checkpoint policy。提交前資料為 663 個實體行且含 8 行串接物件，修復後現行"
    "資料為 695 筆有效 JSON，但因未保存訓練資料 hash、快照或 log，不能斷定實際採用何版。"
    "現行入口預設使用 bf16 LoRA，前代 Qwen 路徑才使用 4-bit QLoRA。因實際 run 之環境快照未"
    "保存，表中可覆寫參數不得解讀為經日誌確認之執行值。"
    )

TRAIN_GENERATION_DESC = (
    "本表區分現行 Gemma 訓練入口預設與推論組態。訓練入口預設採 bf16 LoRA，未使用 BitsAndBytes／NF4。服務"
    "映像以 SDPA 執行注意力，Compose 權重模式預設 FP8 且可覆寫 bf16，惟本輪未保存實際啟動"
    "環境快照。生成採 do_sample=False，2026-07 重放每次請求 max_new_tokens=8192。"
    )
LIM5 = (
    "（5） 統計檢定與 RAG 消融：本研究之比較以各維度平均分數與逐案計數呈現，受限於 44 筆"
    "樣本規模，未進行統計顯著性檢定（如成對樣本之無母數檢定），各表差異之統計穩健性尚待"
    "更大樣本下確認。前次 RAG 重放誤沿用前代 Qwen 嵌入之閾值 0.7，因而 44 案皆為零命中。"
    "改用 EmbeddingGemma 校準閾值 0.32 後，獨立檢索稽核有 25／44 案命中，合計取回 133 份"
    "規則文件，說明原零命中屬設定不相容而非 RAG 必然無效。惟此閾值係於同一 44 案查詢上"
    "校準，且命中不等同規則相關或審查品質提升。服務程式現已支援相關語料檢索、同數量無關"
    "語料檢索、全規則直接注入與關閉 RAG 等啟動模式，並以語料 SHA-256 記錄版本，但這些組別"
    "之 44 案生成層重放尚未完成，人工相關性複核亦仍屬必要。"
)
LIM6 = (
    "（6） 客觀評測基準之性質：§5.2 之客觀問題偵測以 gpt-5.6-sol 僅讀原始碼所列之真問題"
    "作為標註基準，其覆蓋比對亦由 Claude 執行，二者皆為大型語言模型產生之 model-derived "
    "reference，而非人工黃金標準，可能承載模型自身之問題偏好。以人工標註之黃金基準重新量測"
    "召回率與精確率，以及對標註基準之人工複核，均列為未來工作。"
)

FUTURE_REPLAY = (
    "此項工作已先完成現行 Gemma-4-31B-it 加掛 LoRA 轉接器之 44 案重放與產物歸檔。原重放"
    "誤沿用前代 Qwen 嵌入閾值 0.7，因而沒有規則文件命中，不能作為有效 RAG 消融。改採 "
    "EmbeddingGemma 校準閾值 0.32 後，獨立檢索稽核顯示 25／44 案命中，合計取回 133 份規則"
    "文件，完整紀錄存於 datas/Results/2026-07-20-rag-hit-audit.json。此結果只證明檢索設定已"
    "恢復作用。為進一步分離檢索相關性與單純增加上下文的效果，服務程式已新增四種可鎖定之"
    "啟動組別：相關語料檢索、同為 19 份之無關語料檢索、19 份相關規則全量直接注入，以及關閉 "
    "RAG，健康檢查會回報模式、語料名稱與語料 SHA-256。各組仍須以相同模型、提示詞、解碼參數"
    "與 token 預算完成 44 案重放。在該比較完成前，本研究不以檢索命中直接宣稱 RAG 提升審查"
    "品質。"
)


# --------------------------------------------------------------------------- #
def normalise_fonts(doc):
    body = doc.element.body
    has_content = re.compile(r"[A-Za-z0-9一-鿿　-〿＀-￯]")
    fixed = 0
    for r in body.findall(".//" + qn("w:r")):
        text = "".join(t.text or "" for t in r.findall(qn("w:t")))
        if not text or not has_content.search(text):
            continue
        rpr = r.find(qn("w:rPr"))
        rfonts = rpr.find(qn("w:rFonts")) if rpr is not None else None
        if rfonts is not None and rfonts.get(qn("w:ascii")) == "Cambria Math":
            continue
        if rpr is None:
            rpr = r.makeelement(qn("w:rPr"), {})
            r.insert(0, rpr)
        if rfonts is None:
            rfonts = rpr.makeelement(qn("w:rFonts"), {})
            rpr.insert(0, rfonts)
        for slot in ("w:ascii", "w:hAnsi", "w:cs"):
            rfonts.set(qn(slot), "Times New Roman")
        rfonts.set(qn("w:eastAsia"), "標楷體")
        fixed += 1
    print(f"OK fonts normalised on {fixed} runs")


def main():
    doc = Document(SRC)
    parent = doc.paragraphs[0]._parent

    # 1. capture templates + paragraph formats BEFORE any structural change.
    tmpl_2col = copy.deepcopy(doc.tables[1]._element)   # 表二 train params, 2 col
    tmpl_3col = copy.deepcopy(doc.tables[6]._element)   # 表七 old GPT-5, 3 col
    cap_fmt = para_fmt(find_body_para(doc, "表六、採用 CRSCORE++方法的多階段提示詞整體評估",
                                      exact=True))
    desc_fmt = para_fmt(find_body_para(doc, "本表以 CRSCORE++ 評分法在 comprehensiveness"))
    lim_fmt = para_fmt(find_body_para(doc, "統計檢定與模組消融範圍"))
    fig_cap_fmt = para_fmt(find_body_para(doc, "圖一、系統架構圖", exact=True))
    fig_desc_fmt = para_fmt(find_body_para(doc, "本研究所提之程式碼審查框架採分層式架構"))

    # 2. add a separate engineering deployment figure after the concept view.
    insert_figure_block(
        find_body_para(doc, "圖一、系統架構圖", exact=True),
        parent,
        fig_cap_fmt,
        fig_desc_fmt,
        DEPLOYMENT_PNG,
        DEPLOY_CAP,
        DEPLOY_DESC,
    )
    rename_caption(doc, "圖四、LLM-as-a-Judge 流程", "圖五、LLM-as-a-Judge 流程")
    rename_caption(doc, "圖三、程式碼審查流程", "圖四、程式碼審查流程")
    rename_caption(doc, "圖二、訓練流程", "圖三、訓練流程")
    edit_substr(
        doc,
        "code review flow figure reference",
        "本研究之程式碼審查系統整體架構如圖二所示",
        "如圖二所示",
        "如圖四所示",
    )
    edit_substr(
        doc,
        "judge flow figure reference",
        "本研究之 LLM-as-a-Judge 流程如圖四所示",
        "如圖四所示",
        "如圖五所示",
    )

    # 3. insert objective + LLM-judge blocks AFTER the CRSCORE++ description.
    cursor = find_body_para(doc, "本表以 CRSCORE++ 評分法在 comprehensiveness")._p
    cursor = insert_table_block(cursor, parent, tmpl_3col, cap_fmt, desc_fmt,
                                OBJ_CAP, OBJ_ROWS, OBJ_DESC)
    cursor = insert_table_block(cursor, parent, tmpl_3col, cap_fmt, desc_fmt,
                                JUDGE_CAP, JUDGE_ROWS, JUDGE_DESC)

    # 4. insert dataset-statistics block at the end of §4.1.
    ds_anchor = find_body_para(doc, "這 44 筆高品質測試資料也可以作為教師模型與學生模型")
    insert_table_block(ds_anchor._p, parent, tmpl_2col, cap_fmt, desc_fmt,
                       DS_CAP, DS_ROWS, DS_DESC)

    # 5. delete old unsourced tables 表七–表十一.
    for cap in ("表七、不同提示詞設計比較 （GPT-5）",
                "表八、不同提示詞設計比較 （Gemini-3）",
                "表九、提示詞設計比較 （Gemini-3）",
                "表十、微調後與基礎模型比較 （GPT-5）",
                "表十一、微調後與基礎模型比較 （Gemini-3）"):
        delete_table_block(doc, cap)

    # 6. rename body captions (train tables +1, CRSCORE -> 表七, human -> 表十/十一).
    rename_caption(doc, "表二、模型訓練參數配置", "表三、模型訓練參數配置")
    rename_caption(doc, "表三、模型訓練參數說明", "表四、模型訓練參數說明")
    rename_caption(doc, "表四、模型訓練量化與程式碼審查生成參數設定",
                   "表五、模型訓練量化與程式碼審查生成參數設定")
    rename_caption(doc, "表五、模型套件與用途", "表六、模型套件與用途")
    rename_caption(doc, "表六、採用 CRSCORE++方法的多階段提示詞整體評估",
                   "表七、採用 CRSCORE++方法的多階段提示詞整體評估")
    rename_caption(doc, "表十二、不同提示詞設計比較（人工評分）", HUMAN10_CAP)
    rename_caption(doc, "表十三、微調程式碼審查結果 （人工評分）", HUMAN11_CAP)
    replace_table_header_after_caption(doc, parent, HUMAN10_CAP,
                                       ["維度", "多階段提示詞", "單一提示詞"])

    # Replace training tables with the configuration used by the actual Gemma entrypoint.
    replace_table_after_caption(doc, parent, "表三、模型訓練參數配置", TRAIN_ROWS)
    replace_table_after_caption(doc, parent, "表四、模型訓練參數說明", TRAIN_EXPLAIN_ROWS)
    replace_table_after_caption(doc, parent,
                                "表五、模型訓練量化與程式碼審查生成參數設定",
                                TRAIN_GENERATION_ROWS)
    replace_table_after_caption(doc, parent, "表六、模型套件與用途", PACKAGE_ROWS)
    _set_cell(doc.tables[0].rows[-1].cells[2], "✓（現行 Gemma LoRA，前代 Qwen QLoRA）")
    print("OK distinguish LoRA/QLoRA in literature table")

    # 7. prose rewrites / renumbering.
    rewrite_paragraph(doc, "Chinese abstract training scope",
                      "將 LLM 直接用於審查時，會遇到三個彼此牽連之問題",
                      ABSTRACT_METHOD_ZH)
    rewrite_paragraph(doc, "English abstract training scope",
                      "Applying LLMs directly to code review raises three intertwined problems",
                      ABSTRACT_METHOD_EN)
    rewrite_paragraph(doc, "§1.4 problem mapping",
                      "對應地，本研究將三個問題明確映射至三項設計", PROBLEM_MAPPING)
    rewrite_paragraph(doc, "§2.9 distillation evidence",
                      "對本研究的意義：知識蒸餾是本研究把", KD_SIGNIFICANCE)
    rewrite_paragraph(doc, "§4.1 dataset provenance",
                      "為公平地評估審查框架，測試題目本身不能偏心", DATASET_SOURCE)
    rewrite_paragraph(doc, "§4.1 dataset size",
                      "我們產生了 44 筆測試資料", DATASET_SIZE)
    rewrite_paragraph(doc, "§4.1 dataset scope",
                      "這 44 筆高品質測試資料也可以作為教師模型與學生模型", DATASET_SCOPE)
    rewrite_paragraph(doc, "§1.5 training contribution",
                      "2. 以知識蒸餾結合 QLoRA 之輕量化教師–學生訓練流程",
                      "2. 以模型生成回答進行 response-based distillation／SFT，並以 LoRA 作參數高效"
                      "微調。現行 Gemma 訓練入口預設採 bf16 LoRA，前代 Qwen 配置另採 QLoRA。"
                      "作者確認教師模型為 ChatGPT 5.4，惟其參數規模未公開，逐筆生成 provenance "
                      "亦未留存，故本研究不宣稱完成特定大型教師至較小學生之尺寸壓縮。")
    rewrite_paragraph(doc, "§6.2 training contribution",
                      "其二，於模型輕量化上，以知識蒸餾結合 QLoRA",
                      "其二，於訓練工程上，以模型生成回答進行 response-based distillation／SFT，"
                      "並以 LoRA 適配器降低可訓練參數量、保留可拆卸性。現行 Gemma 訓練入口預設採 bf16 LoRA，"
                      "前代 Qwen 配置另採 QLoRA。作者確認教師為 ChatGPT 5.4，惟其參數規模未公開，"
                      "逐筆生成 provenance 亦未留存，故不主張已驗證模型尺寸壓縮。")
    edit_substr(doc, "§1.6 training terminology",
                "本論文共分為六章",
                "知識蒸餾與 QLoRA 微調管線",
                "response-based distillation 與 LoRA 微調管線")
    edit_substr(doc, "§2.10 research gap training terminology",
                "綜合上述比較可知",
                "知識蒸餾＋QLoRA 輕量化微調",
                "response-based distillation＋LoRA 參數高效微調")
    rewrite_paragraph(doc, "§2.8 QLoRA scope",
                      "對本研究的意義：QLoRA 是本研究訓練端之具體實作選擇", QLORA_ROLE)
    rewrite_paragraph(doc, "§3.1 deployment precision",
                      "線上推論以本機後端（LocalHF）以 bf16 載入現行基底", ARCH_DEPLOY)
    rewrite_paragraph(doc, "§3.2.1 actual Gemma training",
                      "本研究的訓練流程採用知識蒸餾與低秩適應微調並行", TRAIN_FLOW)
    rewrite_paragraph(doc, "§4.3.2 current student config",
                      "本研究實驗所用之學生模型及其 QLoRA 微調配置如表二所列",
                      "本研究現行學生模型為 google/gemma-4-31B-it，訓練入口預設採 bf16 LoRA，前代 Qwen "
                      "配置則另採 4-bit QLoRA。現行 Gemma 之可追溯設定如表三所列。")
    rewrite_paragraph(doc, "§4.3.1 split hardware generations",
                      "本研究所有實驗皆在Intel Xeon Gold 6526Y", TRAIN_ENV)
    rewrite_paragraph(doc, "table 3 description",
                      "本表呈現實驗以 QLoRA 微調表列學生模型時所採用", TRAIN_TABLE_DESC)
    rewrite_paragraph(doc, "table 5 description",
                      "本表分為兩部分：上半列出 QLoRA 微調所用之訓練量化配置",
                      TRAIN_GENERATION_DESC)
    edit_substr(doc, "§3.1 RAG threshold",
                "審查管線於各階段以檢索增強生成取得外部知識",
                "相似度閾值 ≥ 0.7",
                "EmbeddingGemma 校準相似度閾值 ≥ 0.32（前代 Qwen 嵌入為 0.7）")
    edit_substr(doc, "§3.4 RAG threshold",
                "為強化審查的依據與一致性，系統建立 RAG 子流程",
                "相似度 threshold ≥ 0.7",
                " EmbeddingGemma 相似度 threshold ≥ 0.32")
    edit_substr(doc, "中文摘要-消融",
                "消融分析作為整合框架內部之貢獻分解，顯示多階段提示詞流程為單項貢獻最大之一環",
                "消融分析作為整合框架內部之貢獻分解，顯示多階段提示詞流程為單項貢獻最大之一環"
                "（+34 分），LoRA 微調之邊際貢獻雖較小（+2 分），卻解決不同問題，在有限資源"
                "下保留並穩定教師之審查能力，與提示詞流程互補而非競爭。",
                "於 44 筆基準之客觀問題偵測評測中，多階段流程較單一提示詞找出更多真實問題"
                "（召回率 0.592 對 0.498、F1 0.740 對 0.663）且精確率相同（皆 0.988），"
                "此為其於問題偵測完整性上之邊際貢獻，而主觀行文品質兩者相當、單一提示詞於"
                "呈現略佳，二者互補而非彼此可取代，RAG 與微調之獨立效益則受本輪合成資料與"
                "部署條件所限未能乾淨切分，本研究因此定位三者為協同互補之整合框架。")
    edit_substr(doc, "英文摘要-消融",
                "As a contribution decomposition within the integrated framework, the ablation shows",
                "As a contribution decomposition within the integrated framework, the ablation "
                "shows that the multi-stage prompting workflow is the single largest contributor "
                "(+34), while LoRA fine-tuning yields a smaller marginal gain (+2) yet solves a "
                "different problem, preserving and stabilizing the teacher's reviewing ability "
                "under limited resources, and thus complements rather than competes with the "
                "prompting workflow.",
                "On an objective bug-detection evaluation over the 44-case benchmark, the "
                "multi-stage workflow finds more real issues than a single prompt (recall 0.592 "
                "vs 0.498, F1 0.740 vs 0.663) at identical precision (0.988 for both), which is "
                "its marginal contribution to issue-detection completeness; on subjective prose "
                "quality the two are comparable, with the single prompt slightly ahead on "
                "presentation, so the two metrics are complementary. The independent effects of "
                "RAG and fine-tuning could not be cleanly isolated on this synthetic benchmark "
                "and deployed model.")
    rewrite_paragraph(doc, "RQ2", "在固定參數規模下，僅引入多階段提示詞而不進行模型微調", RQ2)
    rewrite_paragraph(doc, "RQ3", "多階段提示詞與模型微調二者，對最終分數之邊際貢獻分別為何", RQ3)
    rewrite_paragraph(doc, "§1.2 evidence scope", "這三個問題並非各自獨立，而是構成一條因果鏈",
                      MOTIVATION_REFRAME)
    rewrite_paragraph(doc, "§1.4 ablation scope", "為釐清整合框架內部各元件之相對作用",
                      OBJECTIVE_ABLATION)
    rewrite_paragraph(doc, "§1.5 contribution evidence", "本研究之貢獻，整體可視為一套回應前節",
                      CONTRIBUTION_INTRO)
    rewrite_paragraph(doc, "§2.7 judge provenance", "對本研究的意義：LLM-as-a-Judge", LIT_JUDGE)
    rewrite_paragraph(doc, "§3.3 CoT role", "在整條 CoT 中的角色：本步驟為整條鏈之「推理深化層」",
                      COT_ROLE)
    rewrite_paragraph(doc, "§3.3 CoT caption", "本圖為套用 Chain-of-Thought 提示詞後", COT_CAPTION)
    edit_substr(doc, "§1.5-1",
                "經 §5.2 表七、表八之 LLM 評分消融實驗，本設計相對單一提示詞之邊際貢獻最為顯著。",
                "經 §5.2 表七、表八之 LLM 評分消融實驗，本設計相對單一提示詞之邊際貢獻最為顯著。",
                "經 §5.2 表八之客觀問題偵測評測，本設計相對單一提示詞可找出更多真實問題，於召回率與 F1 上取得提升而精確率不降。")
    rewrite_paragraph(doc, "§4.2 intro", "換言之，表六回答「是否優於既有基準」", S42_INTRO)
    make_para_like(find_body_para(doc, "換言之，表七以 CRSCORE++ 三維度與既有基準作描述性對照")._p,
                   parent, desc_fmt, S42_OBJ_METHOD)
    print("OK insert §4.2 objective-method para")
    edit_substr(doc, "§4.2.4",
                "使表十二、表十三可直接與表七至表十一比較",
                "使表十二、表十三可直接與表七至表十一比較",
                "使表十、表十一可直接與表九比較")
    edit_substr(doc, "§4.2.4-品質依據",
                "並作為表十二與表十三的品質依據",
                "表十二與表十三", "表十與表十一")
    edit_substr(doc, "§4.2.2-正規化表號",
                "表六所列即依 CRSCORE++ 此一正規化尺度呈現",
                "表六所列", "表七所列")
    edit_substr(doc, "§4.2.2-正規化公式",
                "CRSCORE++ 於其原始論文即將 1 至 5 分之評等正規化至 0 至 1 區間",
                "CRSCORE++ 於其原始論文即將 1 至 5 分之評等正規化至 0 至 1 區間（1 分對應 0.2、5 分對應 1.0）",
                "CRSCORE++ 於其原始論文即以正規化分數＝原始評等／5，將 1 至 5 分之評等映射至 0 至 1 區間（1 分對應 0.2、5 分對應 1.0）")
    edit_substr(doc, "§4.2.2-RQ1框架",
                "本研究將逐案分數彙整為表六之平均值",
                "本研究將逐案分數彙整為表六之平均值，因此表六的三維度分數主要用於回答 "
                "RQ1，亦即本研究框架是否在既有 CRSCORE++ 尺度下優於基準。",
                "本研究將逐案分數彙整為表七之平均值，此三維度分數用於與既有 CRSCORE++ "
                "基準作描述性對照，因基準欄之裁判模型與受評資料皆與 Ours 欄不同，此對照"
                "僅屬指示性，不逕以之回答 RQ1。")
    edit_substr(doc, "§4.2.3-五維對照",
                "這五維度不是取代表六",
                "這五維度不是取代表六，而是把表六看不出的細節拆開，使表七至表十一能分析"
                "多階段提示詞、微調與基礎模型之間的差異來源。",
                "這五維度不是取代表七，而是把表七看不出的細節拆開，使表九之 "
                "LLM-as-a-Judge 五維主觀評分與表十、表十一之前代人工評分能進一步分析"
                "多階段提示詞、微調與基礎模型之間的差異來源。")
    edit_substr(doc, "§5.2.3.1-彙整表號",
                "最後再彙整為表十二與表十三",
                "表十二與表十三", "表十與表十一")
    edit_substr(doc, "§4.3.2-表三desc",
                "本表逐項解釋表二所列訓練參數",
                "本表逐項解釋表二所列訓練參數", "本表逐項解釋表三所列訓練參數")
    edit_substr(doc, "§4.3.4-表六",
                "集中規劃並呈現在表五", "呈現在表五", "呈現在表六")
    rewrite_paragraph(doc, "§5.1 intro", "各關鍵設計對生成結果之影響，本節將 CRSCORE++ 三維度", S1_INTRO)
    rewrite_paragraph(doc, "CRSCORE desc", "本表以 CRSCORE++ 評分法在 comprehensiveness", CRSCORE_DESC)
    rewrite_paragraph(doc, "§5.1 metric relation", "透過兩套 LLM-as-a-Judge 評分結果", S51_RELATION)
    rewrite_paragraph(doc, "§5.2.2 heading", "5.2.2 提示詞設計比較 LLM-as-a-Judge評分",
                      "5.2.2 LLM-as-a-Judge 主觀品質比較")
    rewrite_paragraph(doc, "§5.2.2.1 heading", "5.2.2.1 單提示詞與多階提示詞",
                      "5.2.2.1 現行 Gemma 雙裁判比較")
    rewrite_paragraph(doc, "§5.2.2 intro", "在 LLM-as-a-Judge 的評估框架中，提示詞的設計", S522_INTRO)
    rewrite_paragraph(doc, "§5.2.2 design", "在本研究中，評分階段使用三種大型語言模型", S522_DESIGN)
    delete_paragraph(doc, "§5.2.2 old prompt paragraph 1", "為了分析提示詞設計對評分結果的影響")
    delete_paragraph(doc, "§5.2.2 old prompt paragraph 2", "第二種為 本研究提出之多階段程式碼審查提示詞設計")
    rewrite_paragraph(doc, "§5.2.2 result", "在實驗流程中，程式碼審查結果會分別套用不同的評分提示詞",
                      S522_RESULT)
    rewrite_paragraph(doc, "human10 desc", "本表以人工評分為基準，比較本研究多階段提示詞與基礎模型", HUMAN10_DESC)
    rewrite_paragraph(doc, "human11 desc", "本表以人工評分比較有微調", HUMAN11_DESC)
    edit_substr(doc, "§5.2.1.2-表九",
                "因此更適合分析表七至表十一之消融結果",
                "更適合分析表七至表十一之消融結果",
                "更適合分析表九之五維主觀評分與表十、表十一之前代人工評分結果")
    rewrite_paragraph(doc, "§5.2.3 protocol", "人工評分透過附錄一", HUMAN_PROTOCOL)
    rewrite_paragraph(doc, "§5.2.3.2 human calc", "分數彙整時，先對同一受評方法", HUMAN_CALC)
    rewrite_paragraph(doc, "§5.3.1-A", "本節將三張結果表串成一條論證", S531_A)
    rewrite_paragraph(doc, "§5.3.1-B", "由上述一致性檢核可知，兩套獨立評審", S531_B)
    rewrite_paragraph(doc, "§5.3.2-A", "次看整合框架內部各元件之貢獻分解", S532_A)
    rewrite_paragraph(doc, "§5.3.2-B", "此一比例支持兩點觀察", S532_B)
    rewrite_paragraph(doc, "§5.3.2-C", "須特別澄清，+34 與 +2", S532_C)
    rewrite_paragraph(doc, "§5.3.3-A", "末看與基準方法之整體對照", S533_A)
    rewrite_paragraph(doc, "§5.3.3-B", "RQ1 得到正面回答", S533_B)
    rewrite_paragraph(doc, "§6.1 evidence-tier intro", "本研究針對傳統人工程式碼審查所面臨之效率不足", C61_INTRO)
    rewrite_paragraph(doc, "§6.1 conclusion", "經 CRSCORE++ 與 LLM-as-a-Judge-Our 雙重量化評估並由人工評分交叉驗證", C61)
    rewrite_paragraph(doc, "§6.2-1", "本研究之核心貢獻（已於 §5 達成並具量化證據", C62_1)
    rewrite_paragraph(doc, "§6.2-3", "其三，於評估方法上，提出 LLM-as-a-Judge-Our 五項百分制", C62_3)
    rewrite_paragraph(doc, "§6.2 engineering scope", "此外，本研究隨附之開源框架已將上述審查流程整合",
                      CONTRIBUTION_CLOSE)
    rewrite_paragraph(doc, "§6.3-1", "資料規模與多樣性：實驗以 GPT-5 與 Copilot 共同生成 44 筆 Python 測試資料", LIM1)
    rewrite_paragraph(doc, "§6.3-2", "評審模型偏差：評估採 LLM-as-a-Judge 機制", LIM2)
    rewrite_paragraph(doc, "§6.3-3", "微調範圍：本研究之量化評估僅及於表二所列", LIM3)
    rewrite_paragraph(doc, "§6.3-5", "統計檢定與模組消融範圍", LIM5)
    rewrite_paragraph(doc, "§6.4.1 calibrated replay", "此項工作之第一步，以新基座模型重放整條審查管線",
                      FUTURE_REPLAY)
    edit_substr(doc, "appendix",
                "此設計使人工評分可直接對照第五章表十二與表十三",
                "第五章表十二與表十三", "第五章表十與表十一")

    # 8. new §6.3 limitation (6) after the rewritten (5).
    lim6_el = make_para_like(find_body_para(doc, "統計檢定與 RAG 消融")._p, parent, lim_fmt, LIM6)
    print("OK insert §6.3 limitation (6)")
    lim7_el = make_para_like(lim6_el, parent, lim_fmt, LIM7)
    print("OK insert §6.3 limitation (7)")
    make_para_like(lim7_el, parent, lim_fmt, LIM8)
    print("OK insert §6.3 limitation (8)")

    # 9. figure/table cached directories.
    insert_tof_clone_before(doc, "圖二、訓練流程", DEPLOY_CAP)
    set_tof(doc, "圖四、LLM-as-a-Judge 流程", "圖五、LLM-as-a-Judge 流程")
    set_tof(doc, "圖三、程式碼審查流程", "圖四、程式碼審查流程")
    set_tof(doc, "圖二、訓練流程", "圖三、訓練流程")
    insert_tof_clone_before(doc, "表二、模型訓練參數配置", DS_CAP)
    set_tof(doc, "表二、模型訓練參數配置", "表三、模型訓練參數配置")
    set_tof(doc, "表三、模型訓練參數說明", "表四、模型訓練參數說明")
    set_tof(doc, "表四、模型訓練量化與程式碼審查生成參數設定",
            "表五、模型訓練量化與程式碼審查生成參數設定")
    set_tof(doc, "表五、模型套件與用途", "表六、模型套件與用途")
    set_tof(doc, "表六、採用 CRSCORE++方法的多階段提示詞整體評估",
            "表七、採用 CRSCORE++方法的多階段提示詞整體評估")
    set_tof(doc, "表七、不同提示詞設計比較 （GPT-5）", OBJ_CAP)
    set_tof(doc, "表八、不同提示詞設計比較 （Gemini-3）", JUDGE_CAP)
    delete_tof(doc, "表九、提示詞設計比較 （Gemini-3）")
    delete_tof(doc, "表十、微調後與基礎模型比較 （GPT-5）")
    delete_tof(doc, "表十一、微調後與基礎模型比較 （Gemini-3）")
    set_tof(doc, "表十二、不同提示詞設計比較（人工評分）", HUMAN10_CAP)
    set_tof(doc, "表十三、微調程式碼審查結果 （人工評分）", HUMAN11_CAP)

    # 10. normalise fonts (four-slot rFonts; skip math/symbol runs).
    normalise_fonts(doc)

    doc.save(DST)
    print(f"SAVED {DST.name}")


if __name__ == "__main__":
    main()
