"""Apply the thesis drop-in blocks from paper_inserts.md into 論文_v1.8.docx,
saving 論文_v1.9.docx (original untouched).

Each blockquote is parsed into typed lines — section headings (N.M / N.M.K),
bullet items (`- `), numbered / lettered list items (`1. ` / `(a) `) and body
prose — and rendered as Word paragraphs whose paragraph/run formatting is COPIED
from an existing template paragraph of the same level (so the new §3.5/§3.6/§3.7
headings match the live §3.4 / §3.2.1 look, and body text matches §3.4 body).
Headings are additionally bolded for visibility (paper_rule 子章節標題必須可見).
"""
import copy
import re
import sys
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph

sys.stdout.reconfigure(encoding="utf-8")
base = Path(__file__).parent
MD = (base / "paper_inserts.md").read_text(encoding="utf-8").split("\n")

_CJK = r"‘-‟　-〿㐀-鿿＀-￯"
_INLINE = re.compile(r"\*\*(.+?)\*\*|``([^`]+?)``|`([^`]+?)`")


def classify(s):
    if re.match(r"^\d+\.\d+\.\d+\s+\S", s) and len(s) < 60:
        return "h3"
    if re.match(r"^\d+\.\d+\s+\S", s) and len(s) < 60:
        return "h2"
    if s.startswith("- "):
        return "bullet"
    if re.match(r"^\d+\.\s", s):
        return "listnum"
    if re.match(r"^[（(][a-z0-9]+[）)]", s):
        return "listalpha"
    return "body"


def merge(lines):
    text = " ".join(lines).replace("\\ ", "")
    return re.sub(rf"(?<=[{_CJK}])\s+(?=[{_CJK}])", "", text)


class _TypedCollector:
    """Accumulates blockquote content lines into (kind, merged-text) items."""

    def __init__(self):
        self.items = []
        self._cur = []
        self._cur_kind = None
        self._in_content = False

    def flush(self):
        if self._cur:
            self.items.append((self._cur_kind, merge(self._cur)))
            self._cur, self._cur_kind = [], None

    def feed(self, line):
        if line.startswith("**內容**") or line.startswith("**標題與內容**"):
            self._in_content = True
            return
        if not self._in_content:
            return
        if not line.startswith(">"):
            if line.strip() == "":
                self.flush()
            return
        self._feed_quoted(line[1:].strip())

    def _feed_quoted(self, content):
        if content == "":
            self.flush()
            return
        kind = classify(content)
        # A heading / list item always starts a new item; body prose only
        # starts one when nothing is being accumulated yet.
        if kind != "body" or self._cur_kind is None:
            self.flush()
            self._cur_kind, self._cur = kind, [content]
        else:
            self._cur.append(content)


def extract_typed(header_prefix):
    start = next(i for i, line in enumerate(MD) if line.startswith(header_prefix))
    collector = _TypedCollector()
    for line in MD[start + 1:]:
        if line.startswith("### ") or line.rstrip() == "---":
            break
        collector.feed(line)
    collector.flush()
    return collector.items


def parse_inline(text):
    out, pos = [], 0
    for m in _INLINE.finditer(text):
        if m.start() > pos:
            out.append((text[pos:m.start()], False))
        if m.group(1) is not None:
            out.append((m.group(1), True))
        else:
            out.append((m.group(2) or m.group(3), False))
        pos = m.end()
    if pos < len(text):
        out.append((text[pos:], False))
    return out or [(text, False)]


def find(doc, sub, exact=False):
    for p in doc.paragraphs:
        if (p.text.strip() == sub) if exact else (sub in p.text):
            return p
    raise SystemExit(f"TEMPLATE/ANCHOR NOT FOUND: {sub}")


def fmt_of(p):
    ppr = p._p.find(qn("w:pPr"))
    rpr = None
    for r in p.runs:
        if r.text.strip():
            rpr = r._element.find(qn("w:rPr"))
            break
    return ppr, rpr


def _para_style(kind, T):
    if kind == "h2":
        return (*T["h2"], True)
    if kind == "h3":
        return (*T["h3"], True)
    return (*T["body"], False)


def _para_text(kind, text):
    if kind in ("h2", "h3"):  # normalise "3.5  標題" → "3.5 標題" (match the doc's one-space headings)
        return re.sub(r"^(\d+(?:\.\d+)+)\s{2,}", r"\1 ", text, count=1)
    if kind == "bullet":
        return "‧ " + (text[2:] if text.startswith("- ") else text)
    return text


def make_para(after_elem, parent, kind, text, T):
    ppr, rpr, bold = _para_style(kind, T)
    text = _para_text(kind, text)
    new_p = OxmlElement("w:p")
    after_elem.addnext(new_p)
    if ppr is not None:
        new_p.append(copy.deepcopy(ppr))
    para = Paragraph(new_p, parent)
    for seg, b in parse_inline(text):
        run = para.add_run(seg)
        if rpr is not None:
            run._element.insert(0, copy.deepcopy(rpr))
        if b or bold:
            run.bold = True
    return new_p


def apply_block(header, after_elem, parent, T):
    cursor = after_elem
    n = 0
    for kind, text in extract_typed(header):
        cursor = make_para(cursor, parent, kind, text, T)
        n += 1
    print(f"OK {header.strip()}: +{n} paragraphs")
    return cursor


doc = Document(base / "論文_v1.8.docx")
parent = doc.paragraphs[0]._parent
T = {
    "body": fmt_of(find(doc, "AI 程式碼審查與程式碼編輯環境整合能讓")),   # §3.4 body, 14pt
    "h2": fmt_of(find(doc, "3.4 與程式碼編輯環境整合", exact=True)),       # §3.x heading, 20pt
    "h3": fmt_of(find(doc, "3.2.1 多階提示詞", exact=True)),               # §3.x.y heading, 18pt
}

# §3.5 / §3.6 / §3.7 chained after §3.4 body (before 第四章).
cursor = find(doc, "AI 程式碼審查與程式碼編輯環境整合能讓")._p
for header in ("### 2.2 ", "### 2.3 ", "### 2.4 "):
    cursor = apply_block(header, cursor, parent, T)

# §5.3 after §5.2 tail (before 第六章).
cursor = find(doc, "本區塊為人工評分用的評分標準說明")._p
apply_block("### 2.5 ", cursor, parent, T)


def delete_after_until(anchor_p, stop_pred):
    """Delete every paragraph after anchor_p up to (not incl.) the stop match."""
    el = anchor_p._p.getnext()
    while el is not None:
        if stop_pred(Paragraph(el, parent)):
            break
        nxt = el.getnext()
        el.getparent().remove(el)
        el = nxt


def delete_inclusive_until(start_p, stop_pred):
    """Delete start_p and following paragraphs up to (not incl.) the stop match;
    return the element BEFORE start_p so new content can be inserted there."""
    prev = start_p._p.getprevious()
    el = start_p._p
    while el is not None:
        if stop_pred(Paragraph(el, parent)):
            break
        nxt = el.getnext()
        el.getparent().remove(el)
        el = nxt
    return prev


# §1.5 — replace the original four-point list with the seven-contribution list.
h15 = find(doc, "1.5 研究貢獻", exact=True)
delete_after_until(h15, lambda p: p.text.strip().startswith("1.6"))
apply_block("### 2.1 ", h15._p, parent, T)

# §6.4 — replace the whole existing §6.4 (heading + prose) with the structured
# 6.4.1–6.4.5 version (the block carries its own "6.4 未來工作" heading).
h64 = find(doc, "6.4 未來工作", exact=True)
prev64 = delete_inclusive_until(h64, lambda p: "參考文獻" in p.text)
apply_block("### 2.6 ", prev64, parent, T)

doc.save(base / "論文_v1.9.docx")
print("SAVED 論文_v1.9.docx")
