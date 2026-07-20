"""Produce 論文_v3.8.docx from 論文_v3.7.docx (original untouched).

Insert a single-page 致謝 (Acknowledgments) section between the cover page and
the 摘要 front matter:

    cover page → [page break] → 致謝 page → [page break] → 摘要 …

Structure inserted right after the cover's last paragraph ("June, 2026"):
1. A 致謝 heading (Heading 1 style "1", centred, bold-off — mirrors the 摘要
   heading) carrying `<w:pageBreakBefore/>` so it starts on a fresh page after
   the cover.
2. Three short acknowledgment body paragraphs (Normal, 14pt) thanking the
   advisor, the oral-defence committee, and the lab colleagues. Generic and
   truthful — no fabricated names, lab names, events, or institutions.
3. A manual page break (`<w:br w:type="page"/>`) appended after the last body
   paragraph so the original 摘要 starts on the next page.

Every new run sets the four rFonts slots explicitly (ascii/hAnsi/cs=Times New
Roman, eastAsia=標楷體) so no new「含中文卻非標楷體」/「含英文卻設標楷體」run is
introduced. Punctuation is full-width, no dashes, no full-width semicolons.
Usage: python _add_acknowledgments.py
"""
import sys
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

sys.stdout.reconfigure(encoding="utf-8")
base = Path(__file__).parent
SRC = base / "論文_v3.7.docx"
DST = base / "論文_v3.8.docx"

ANCHOR = "June, 2026"
TITLE = "致謝"
BODY = [
    "本論文能順利完成，要向指導教授獻上最誠摯的感謝。感謝老師在研究主題的選定、"
    "方法的規劃與論文的撰寫上悉心指導，並在我遭遇困難與迷惘時耐心點撥、適時鼓勵，"
    "使我得以逐步釐清問題、完成這份研究。老師嚴謹的治學態度，將是我日後持續學習的"
    "榜樣。",
    "感謝各位口試委員在百忙之中撥冗審閱本論文。委員們於口試過程中提出許多寶貴而"
    "深入的意見，指出本研究尚待補強之處並給予具體的修改建議，讓論文的論述與內容"
    "更臻嚴謹完整，使我受益良多，謹此致上由衷的謝意。",
    "也要感謝實驗室的每一位夥伴。在這段研究的日子裡，感謝大家在課業上的相互切磋"
    "與討論，以及在生活上的彼此扶持與陪伴，讓漫長的研究過程多了許多溫暖與力量。"
    "謹將這份成果，與所有曾經幫助、支持與陪伴過我的人分享。",
]


def make_rfonts():
    """A rFonts element with all four slots set (TNR + 標楷體)."""
    rf = OxmlElement("w:rFonts")
    for slot in ("w:ascii", "w:hAnsi", "w:cs"):
        rf.set(qn(slot), "Times New Roman")
    rf.set(qn("w:eastAsia"), "標楷體")
    return rf


def make_sz(val):
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), val)
    szcs = OxmlElement("w:szCs")
    szcs.set(qn("w:val"), val)
    return sz, szcs


def make_bold_off(rpr):
    b = OxmlElement("w:b")
    b.set(qn("w:val"), "0")
    bcs = OxmlElement("w:bCs")
    bcs.set(qn("w:val"), "0")
    rpr.append(b)
    rpr.append(bcs)


def heading_para(text):
    """Heading 1, centred, bold-off, page-break-before — mirrors 摘要."""
    p = OxmlElement("w:p")
    ppr = OxmlElement("w:pPr")
    pstyle = OxmlElement("w:pStyle")
    pstyle.set(qn("w:val"), "1")
    ppr.append(pstyle)
    ppr.append(OxmlElement("w:pageBreakBefore"))
    jc = OxmlElement("w:jc")
    jc.set(qn("w:val"), "center")
    ppr.append(jc)
    prpr = OxmlElement("w:rPr")
    prpr.append(make_rfonts())
    make_bold_off(prpr)
    ppr.append(prpr)
    p.append(ppr)

    r = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    rpr.append(make_rfonts())
    make_bold_off(rpr)
    r.append(rpr)
    t = OxmlElement("w:t")
    t.text = text
    r.append(t)
    p.append(r)
    return p


def body_para(text):
    """Normal body paragraph, 14pt (sz=28)."""
    p = OxmlElement("w:p")
    ppr = OxmlElement("w:pPr")
    prpr = OxmlElement("w:rPr")
    prpr.append(make_rfonts())
    for el in make_sz("28"):
        prpr.append(el)
    ppr.append(prpr)
    p.append(ppr)

    r = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    rpr.append(make_rfonts())
    for el in make_sz("28"):
        rpr.append(el)
    r.append(rpr)
    t = OxmlElement("w:t")
    t.set(qn("xml:space"), "preserve")
    t.text = text
    r.append(t)
    p.append(r)
    return p


def append_page_break(p):
    """Append a manual page-break run (Ctrl+Enter equivalent)."""
    r = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    rpr.append(make_rfonts())
    for el in make_sz("28"):
        rpr.append(el)
    r.append(rpr)
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    r.append(br)
    p.append(r)


doc = Document(SRC)

hits = [p for p in doc.paragraphs if p.text.strip() == ANCHOR]
if len(hits) != 1:
    raise SystemExit(f"ANCHOR NOT UNIQUE ({len(hits)} hits): {ANCHOR}")
anchor = hits[0]._p

# Guard: 致謝 must not already exist (idempotency / no duplicate insertion).
if any(p.text.strip() == TITLE for p in doc.paragraphs):
    raise SystemExit("致謝 heading already present — aborting to avoid duplicate")

new_paras = [heading_para(TITLE)] + [body_para(t) for t in BODY]
append_page_break(new_paras[-1])

cursor = anchor
for el in new_paras:
    cursor.addnext(el)
    cursor = el

doc.save(DST)
print(f"SAVED {DST.name} (inserted 致謝 heading + {len(BODY)} body paras + page break)")
