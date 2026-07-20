"""Produce 論文_v3.9.docx from 論文_v3.8.docx (original untouched).

Data-independent revisions only (no experimental result is added or changed).
Each item is traceable to the cited source (spec: paper/_revision_v3.9_spec.md):

  D5  cover year 一一六 -> 一一五 (matches "June, 2026").
  D3  §1.5 「下列三項」-> 「下列四項」 (items 4-7 are four design contributions).
  D2  TOC 3.3.9 -> 3.5 gap: give body §3.4 an outlineLvl + bookmark, and inject
      the missing cached TOC entry "3.4 程式碼審查流程".
  D4  formalise informal phrasing: drop the「直覺先行：」/「所以呢：」labels
      (keep the plain-language intuition), 塞進提示詞 / 考自己擅長之題 recast.
  B2/B3 §2.9 + §3.2.1: precise response-based KD framing; think field is in
      neither the input nor the loss (source: codes/train/*.py: build_prompt uses
      instruction+question only, label mask = [-100]*prompt_len + answer tokens).
  B1  表四 split into 訓練量化配置 (QLoRA) + 推論配置 (greedy, no temperature/top_p;
      source: codes/util/hf_model_util.py).
  A4  CRSCORE++ 1-5 -> raw/5 normalisation note (source: score.md "1=0.2"), and
      completeness -> comprehensiveness terminology alignment (abstract only).
  C2  primary term = 多階段結構化審查工作流; CoT is one prompting strategy;
      intermediate rationale is not faithful-interpretability evidence.

Out of scope (reserved for the data pass): 表七-表十三, +34/+2, §5.3.2/§5.3.3
([0473-0478]), 表六 fairness/RQ1 wording, 結論三層化.

Usage: python _rewrite_v3_9.py
"""
import copy
import sys
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

sys.stdout.reconfigure(encoding="utf-8")
base = Path(__file__).parent
SRC = base / "論文_v3.8.docx"
DST = base / "論文_v3.9.docx"


def set_run_fonts(run_el):
    rpr = run_el.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    for slot in ("w:ascii", "w:hAnsi", "w:cs"):
        rfonts.set(qn(slot), "Times New Roman")
    rfonts.set(qn("w:eastAsia"), "標楷體")


def find_paragraph(doc, anchor):
    hits = [p for p in doc.paragraphs if anchor in p.text]
    if len(hits) != 1:
        raise SystemExit(f"ANCHOR NOT UNIQUE ({len(hits)}): {anchor[:30]}")
    return hits[0]


def rebuild(p, new_text):
    """Collapse the paragraph to a single run carrying the original rPr plus
    the four explicit rFonts slots (safe for plain prose paragraphs)."""
    src_rpr = None
    for r in p.runs:
        if r.text.strip():
            src_rpr = r._element.find(qn("w:rPr"))
            break
    for r in list(p.runs):
        r._element.getparent().remove(r._element)
    run = p.add_run(new_text)
    if src_rpr is not None:
        run._element.insert(0, copy.deepcopy(src_rpr))
    set_run_fonts(run._element)


def edit_paragraph(doc, label, anchor, repls):
    p = find_paragraph(doc, anchor)
    txt = p.text
    for old, new in repls:
        n = txt.count(old)
        if n != 1:
            raise SystemExit(f"OLD NOT UNIQUE ({n}) [{label}]: {old[:40]}")
        txt = txt.replace(old, new)
    rebuild(p, txt)
    print(f"OK {label}")


# ---------------------------------------------------------------------------
# 1. Plain-paragraph text edits (anchor -> [(old, new), ...])
# ---------------------------------------------------------------------------
PARA_EDITS = [
    ("D5 封面年份", "一一六 年", [("一一六", "一一五")]),
    ("D3 §1.5 貢獻計數", "下列三項屬於本研究隨附之開源框架",
     [("下列三項", "下列四項")]),
    ("D4 §1.2 塞進提示詞", "把整段審查塞進一個提示詞，正是輸出不穩定",
     [("把整段審查塞進一個提示詞", "將全部審查任務置於單一提示詞")]),
    ("D4 §2.5 RAG 塞進", "找出來塞進提示詞，讓模型照這些資料作答",
     [("找出來塞進提示詞", "找出並注入提示詞")]),
    # --- 直覺先行 label removal (keep the intuition content) ---
    ("D4 直覺 §3.3.1", "直覺先行：與其要模型一口氣把整段審查做完",
     [("直覺先行：與其要模型一口氣把整段審查做完",
       "就審查流程之設計而言，與其要模型一口氣把整段審查做完")]),
    ("D4 直覺 §3.3.2", "直覺先行：這一步要找的不是會立刻當掉",
     [("直覺先行：這一步要找的不是", "就本步驟之目的而言，所欲偵測者並非")]),
    ("D4 直覺 §3.3.3", "直覺先行：審查要由淺入深",
     [("直覺先行：審查要由淺入深", "本步驟之設計理念為審查由淺入深")]),
    ("D4 直覺 §3.3.4", "直覺先行：在細看程式碼前",
     [("直覺先行：在細看程式碼前", "就審查之先後次序而言，在細看程式碼前")]),
    ("D4 直覺 §3.3.5", "直覺先行：本步驟是上一步摘要設計",
     [("直覺先行：本步驟是", "本步驟為")]),
    ("D4 直覺 §3.3.6", "直覺先行：這一步要模型「只當檢查器",
     [("直覺先行：這一步要模型", "本步驟要求模型")]),
    ("D4 直覺 §3.3.7", "直覺先行：前面幾步已找出",
     [("直覺先行：前面幾步已找出「有哪些問題」，這一步要回答",
       "前面幾步已找出「有哪些問題」，本步驟則回答")]),
    ("D4 直覺 §3.3.8", "直覺先行：審查到最後，需要有人把",
     [("直覺先行：審查到最後", "就審查之最終彙整而言，審查到最後")]),
    ("D4 直覺 §3.3.9", "直覺先行：要讓每一份審查都按同一套標準走",
     [("直覺先行：要讓每一份審查", "為使每一份審查")]),
    ("D4 直覺 §3.5", "直覺先行：審查若能在開發者寫程式之當下",
     [("直覺先行：審查若能在開發者寫程式之當下",
       "就整合之時機而言，審查若能在開發者寫程式之當下")]),
    ("D4 直覺 §4.1 + 考自己擅長", "直覺先行：要公平地評估審查框架",
     [("直覺先行：要公平地評估審查框架", "為公平地評估審查框架"),
      ("也不能全由單一模型生成而讓受測者「考自己擅長之題」",
       "也不能全由單一模型生成，以避免受測模型評測其自身偏好之題型")]),
    # --- D4 所以呢 §5.3.1 only (§5.3.2/§5.3.3 reserved, left untouched) ---
    ("D4 所以呢 §5.3.1", "所以呢：兩套獨立評審",
     [("所以呢：兩套獨立評審", "由上述一致性檢核可知，兩套獨立評審")]),
    # --- A4 terminology: abstract completeness -> comprehensiveness ---
    ("A4 中文摘要 completeness", "消融分析作為整合框架內部之貢獻分解",
     [("completeness 皆 1.00", "comprehensiveness 皆 1.00")]),
    ("A4 英文摘要 completeness",
     "comparable across all three dimensions (completeness 1.00 for both",
     [("completeness 1.00 for both", "comprehensiveness 1.00 for both")]),
    # --- A4 normalisation note (source: score.md "normalize 0~1,1=0.2") ---
    ("A4 正規化 §4.2.2", "最後於三個維度各給 1–5 分。Comprehensiveness 表示",
     [("最後於三個維度各給 1–5 分。Comprehensiveness 表示",
       "最後於三個維度各給 1–5 分。CRSCORE++ 於其原始論文即將 1 至 5 分之評等"
       "正規化至 0 至 1 區間（1 分對應 0.2、5 分對應 1.0），"
       "表六所列即依 CRSCORE++ 此一正規化尺度呈現。Comprehensiveness 表示")]),
    # --- B2 KD framing (§2.9) ---
    ("B2 §2.9 KD 定義", "使其輸出結果盡可能接近教師模型所生成的內容",
     [("讓學生模型透過微調更新參數，使其輸出結果盡可能接近教師模型所生成的內容。",
       "讓學生模型透過微調更新參數，使其輸出結果盡可能接近教師模型所生成的內容。"
       "須明確界定者，本研究採用之知識蒸餾屬回應式蒸餾"
       "（response-based distillation），即指令蒸餾：學生模型僅以教師生成之回應"
       "（即答案）作為監督訊號進行 SFT，並未取用教師模型之 logits、soft target "
       "或以溫度調節之蒸餾損失，故與前述 Hinton 等人之 soft-target 蒸餾有別。")]),
    # --- B3 think not in loss (§2.9) + C2 traceability framing ---
    ("B3 §2.9 think 未納入 loss",
     "使學生模型學到的是「如何逐步推理地審查」",
     [("本研究以具 CoT 之少樣本提示向教師大模型蒸餾出帶推理軌跡之資料"
       "（而非僅有結論），使學生模型學到的是「如何逐步推理地審查」，"
       "而非死記答案，這正是審查可解釋性得以被保留之原因。",
       "本研究以具 CoT 之少樣本提示向教師大模型生成帶推理軌跡與最終審查答案之資料，"
       "但訓練時學生模型僅以教師之最終審查答案（即資料集之 answer 欄）"
       "作為監督訊號進行 SFT，訓練目標僅涵蓋 answer 欄，think 欄雖存於資料集"
       "卻未納入損失（詳見 §3.2.1）。因此本研究不宣稱學生模型透過訓練習得教師之"
       "推理軌跡，多階段審查之「逐步可追溯」係來自推論時之管線結構"
       "（§3.3 多階段提示詞逐步落檔），而非權重內化之推理。")]),
    # --- B3 §3.2.1 build_prompt(instruction, question); target=answer ---
    ("B3 §3.2.1 build_prompt", "確保模型只學習回答而非提示",
     [("確保模型只學習回答而非提示。",
       "確保模型只學習回答而非提示。訓練輸入模板為 "
       "build_prompt(instruction, question)，訓練目標為 answer 欄，"
       "think 欄雖為資料集之一欄，但既不參與提示組建亦不計入損失。")]),
    # --- C2 workflow framing + interpretability caveat (§3.3.1) ---
    ("C2 §3.3.1 workflow", "在整條 CoT 中的角色：本節之多階提示詞是整個審查鏈",
     [("這正是本研究回應「輸出不穩定」問題之結構性手段。",
       "這正是本研究回應「輸出不穩定」問題之結構性手段。更精確地說，"
       "本研究之方法宜稱為多階段結構化審查工作流"
       "（multi-stage structured reasoning workflow），CoT 僅為其中一種提示策略。"
       "須留意的是，各階段輸出之文字理由並不必然反映模型真實之內部推理，"
       "亦不足以作為模型忠實可解釋性之直接證據，本框架之可追溯性係來自"
       "各階段產物之落檔與流程結構本身，而非模型權重內化推理之忠實映射。")]),
    # --- B1 表四 說明文字 (temperature/top_p correction) ---
    ("B1 表四 說明文字", "本表列出推論階段產生程式碼審查結果時所使用的生成參數",
     [("本表列出推論階段產生程式碼審查結果時所使用的生成參數"
       "（如 temperature、top_p、max_new_tokens 等），"
       "用以確保不同模型在相同生成條件下進行公平比較。",
       "本表分為兩部分：上半列出 QLoRA 微調所用之訓練量化配置"
       "（載入 4-bit 權重、雙重量化、NF4 量化型別與 bf16 計算精度，"
       "來源見 §3.2.1 訓練流程），下半列出推論階段產生程式碼審查結果時之生成配置。"
       "推論以 bf16 載入權重並採 SDPA 注意力，解碼採貪婪策略（do_sample=False），"
       "未使用 temperature／top_p 取樣，亦未設定 repetition_penalty，"
       "以確保不同模型在相同且可重現之生成條件下進行公平比較。")]),
    # --- B1 表四 body caption rename (covers training + inference now) ---
    ("B1 表四 標題", "表四、模型程式碼審查生成參數設定",  # body caption only
     []),  # handled specially below (TOF entry shares this text)
]


def apply_para_edits(doc):
    for label, anchor, repls in PARA_EDITS:
        if not repls:
            continue
        edit_paragraph(doc, label, anchor, repls)


# ---------------------------------------------------------------------------
# 2. 表四 caption rename (body caption via rebuild; TOF entry via run edit)
# ---------------------------------------------------------------------------
OLD_CAP = "表四、模型程式碼審查生成參數設定"
NEW_CAP = "表四、模型訓練量化與程式碼審查生成參數設定"


def rename_table_four_caption(doc):
    body_cap = [p for p in doc.paragraphs
                if p.text.strip() == OLD_CAP
                and "table of figures" not in (p.style.name or "").lower()]
    if len(body_cap) != 1:
        raise SystemExit(f"表四 body caption not unique: {len(body_cap)}")
    rebuild(body_cap[0], NEW_CAP)
    print("OK B1 表四 body caption rename")

    tof = [p for p in doc.paragraphs
           if OLD_CAP in p.text
           and "table of figures" in (p.style.name or "").lower()]
    if len(tof) != 1:
        raise SystemExit(f"表四 TOF entry not unique: {len(tof)}")
    changed = 0
    for t in tof[0]._p.iter(qn("w:t")):
        if t.text and OLD_CAP in t.text:
            t.text = t.text.replace(OLD_CAP, NEW_CAP)
            changed += 1
    if changed != 1:
        raise SystemExit(f"表四 TOF run replace count={changed}")
    print("OK B1 表四 TOF entry rename")


# ---------------------------------------------------------------------------
# 3. 表四 table restructure: 訓練量化配置 (QLoRA) + 推論配置
# ---------------------------------------------------------------------------
TABLE4_ROWS = [
    ("訓練量化配置（QLoRA）", ""),
    ("load_in_4bit", "True"),
    ("bnb_4bit_use_double_quant", "True"),
    ("bnb_4bit_quant_type", "nf4"),
    ("bnb_4bit_compute_dtype", "torch.bfloat16"),
    ("推論配置", ""),
    ("權重載入精度", "bf16"),
    ("注意力實作", "SDPA"),
    ("do_sample", "False（貪婪解碼）"),
    ("temperature", "None"),
    ("top_p", "None"),
    ("max_new_tokens", "32768"),
    ("repetition_penalty", "未設定（無）"),
]


def _set_cell(cell, text):
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    for extra in cell.paragraphs[1:]:
        extra._p.getparent().remove(extra._p)
    p = cell.paragraphs[0]
    for r in list(p.runs):
        r._element.getparent().remove(r._element)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_run_fonts(run._element)
    run.font.size = Pt(14)


def restructure_table_four(doc):
    t = doc.tables[3]
    if [c.text for c in t.rows[0].cells] != ["參數名", "設定"]:
        raise SystemExit("表四 header mismatch; wrong table index")
    if t.rows[1].cells[0].text != "load_in_4bit":
        raise SystemExit("表四 first data row mismatch")
    n_existing = len(t.rows) - 1  # data rows currently present
    # overwrite existing data rows, add rows for the remainder
    for i, (c0, c1) in enumerate(TABLE4_ROWS):
        if i < n_existing:
            row = t.rows[i + 1]
        else:
            row = t.add_row()
        _set_cell(row.cells[0], c0)
        _set_cell(row.cells[1], c1)
    print(f"OK B1 表四 restructured ({len(TABLE4_ROWS)} data rows)")


# ---------------------------------------------------------------------------
# 4. D2 TOC: outlineLvl + bookmark on body §3.4, inject cached TOC entry
# ---------------------------------------------------------------------------
BM_NAME = "_Toc900003400"
BM_ID = "900"


def fix_body_heading_34(doc):
    hits = [p for p in doc.paragraphs if p.text.strip() == "3.4 程式碼審查流程"]
    if len(hits) != 1:
        raise SystemExit(f"body §3.4 heading not unique: {len(hits)}")
    p = hits[0]._p
    pPr = p.find(qn("w:pPr"))
    if pPr is None:
        raise SystemExit("§3.4 heading has no pPr")
    # add outlineLvl=1 (matches sibling §3.5) before rPr so Word regen picks it up
    if pPr.find(qn("w:outlineLvl")) is None:
        ol = OxmlElement("w:outlineLvl")
        ol.set(qn("w:val"), "1")
        rpr = pPr.find(qn("w:rPr"))
        if rpr is not None:
            rpr.addprevious(ol)
        else:
            pPr.append(ol)
    # wrap heading text in a _Toc bookmark for the injected TOC hyperlink/PAGEREF
    if not any(b.get(qn("w:name")) == BM_NAME
               for b in p.iter(qn("w:bookmarkStart"))):
        bs = OxmlElement("w:bookmarkStart")
        bs.set(qn("w:id"), BM_ID)
        bs.set(qn("w:name"), BM_NAME)
        be = OxmlElement("w:bookmarkEnd")
        be.set(qn("w:id"), BM_ID)
        pPr.addnext(bs)
        p.append(be)
    print("OK D2 §3.4 heading outlineLvl + bookmark")


def inject_toc_entry_34(doc):
    sdt = doc.element.body.find(qn("w:sdt"))
    if sdt is None:
        raise SystemExit("TOC sdt not found")
    content = sdt.find(qn("w:sdtContent"))

    def ftext(el):
        return "".join(t.text or "" for t in el.iter(qn("w:t")))

    p35 = None
    for p in content.findall(qn("w:p")):
        if ftext(p).strip().startswith("3.5 與程式碼編輯環境整合"):
            p35 = p
            break
    if p35 is None:
        raise SystemExit("cached TOC 3.5 entry not found")
    # do not double-insert if a 3.4 cached entry already exists
    for p in content.findall(qn("w:p")):
        if ftext(p).strip().startswith("3.4 程式碼審查流程"):
            print("SKIP D2 cached TOC 3.4 already present")
            return
    clone = copy.deepcopy(p35)
    t_changes = {"3.5 ": "3.4 ", "與程式碼編輯環境整合": "程式碼審查流程",
                 "48": "46"}
    seen = {k: 0 for k in t_changes}
    for t in clone.iter(qn("w:t")):
        if t.text in t_changes:
            seen[t.text] += 1
            t.text = t_changes[t.text]
    if any(v != 1 for v in seen.values()):
        raise SystemExit(f"TOC clone text replace counts wrong: {seen}")
    ic = 0
    for it in clone.iter(qn("w:instrText")):
        if it.text and "_Toc232644810" in it.text:
            it.text = it.text.replace("_Toc232644810", BM_NAME)
            ic += 1
    hc = 0
    for h in clone.iter(qn("w:hyperlink")):
        if h.get(qn("w:anchor")) == "_Toc232644810":
            h.set(qn("w:anchor"), BM_NAME)
            hc += 1
    if ic != 1 or hc != 1:
        raise SystemExit(f"TOC clone anchor replace counts wrong: instr={ic} hyp={hc}")
    p35.addprevious(clone)
    print("OK D2 cached TOC 3.4 entry injected (page 46, Word recomputes on update)")


# ---------------------------------------------------------------------------
def main():
    doc = Document(SRC)
    apply_para_edits(doc)
    rename_table_four_caption(doc)
    restructure_table_four(doc)
    fix_body_heading_34(doc)
    inject_toc_entry_34(doc)
    doc.save(DST)
    print(f"SAVED {DST.name}")


if __name__ == "__main__":
    main()
