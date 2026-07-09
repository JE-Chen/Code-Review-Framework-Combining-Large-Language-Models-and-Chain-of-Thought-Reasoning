"""Produce 論文_v3.6.docx from 論文_v3.5.docx (original untouched).

Three audited fixes, nothing else:
1. Table-six judge-inconsistency disclosure (zh abstract, en abstract,
   §5.1 body, table-six caption note) — the three Ours columns are scored
   by Claude with the standard CRSCORE++ Judge Prompt, while the CRSCORE++
   baseline column keeps its originally reported GPT-4o-mini scores; the
   cross-judge comparison is disclosed as indicative only. The caption
   paragraph's now-contradictory「與同一 judge」clause is dropped.
2. SARIF duplicate gloss (§3.8 second occurrence de-glossed).
3. MCP duplicate expansion (§3.5 second occurrence collapsed to MCP;
   §3.7.2 third gloss de-glossed — §1.5 first gloss covers both).

Each edit locates its paragraph by a unique anchor substring, verifies the
old text occurs exactly once, then rebuilds the paragraph as a single run
whose rPr is copied from the original first text run with the four rFonts
slots set explicitly (ascii/hAnsi/cs=Times New Roman, eastAsia=標楷體).
Usage: python _rewrite_v3_6.py
"""
import copy
import sys
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn

sys.stdout.reconfigure(encoding="utf-8")
base = Path(__file__).parent
SRC = base / "論文_v3.5.docx"
DST = base / "論文_v3.6.docx"

# (label, paragraph-anchor, old-substring, new-substring)
EDITS = [
    (
        "1a 中文摘要",
        "實驗於 CRSCORE++ 指標上，以同一 judge",
        "均高於 CRSCORE++ 基準（0.67／0.57／0.63）。",
        "亦高於 CRSCORE++ 基準之原始報告分數（0.67／0.57／0.63，該欄由"
        " GPT-4o-mini 評分，跨 judge 之比較僅屬指示性參考）。",
    ),
    (
        "1b 英文摘要",
        "On the CRSCORE++ metric, under a single judge",
        ", and both are above the CRSCORE++ baseline (0.67/0.57/0.63).",
        ", also exceeding the originally reported CRSCORE++ baseline"
        " (0.67/0.57/0.63, scored by GPT-4o-mini; the cross-judge"
        " comparison is indicative only).",
    ),
    (
        "1c §5.1 正文",
        "末看與基準方法之整體對照",
        "三個 Ours 欄於三維度均高於 CRSCORE++ 基準（0.67／0.57／0.63）。",
        "三個 Ours 欄於三維度亦高於 CRSCORE++ 基準之原始報告分數"
        "（0.67／0.57／0.63）。需說明的是，基準欄沿用 CRSCORE++ 原始評測之"
        " GPT-4o-mini 分數，與三個 Ours 欄之 Claude 評分屬不同 judge，"
        "跨欄比較僅屬指示性參考，三個 Ours 欄彼此間則為同一 judge "
        "下之可比分數。",
    ),
    (
        "1d-1 表六說明：加入 judge 揭露句",
        "本表以 CRSCORE++ 評分法在 comprehensiveness",
        "比較本研究三個學生模型與 CRSCORE++ 基準之整體表現。",
        "比較本研究三個學生模型與 CRSCORE++ 基準之整體表現。"
        "表中三個 Ours 欄均由 Claude 以標準 CRSCORE++ 評分提示詞評分，"
        "CRSCORE++ 基準欄沿用其原始評測之 GPT-4o-mini 分數，跨 judge "
        "之欄間比較僅屬指示性。",
    ),
    (
        # 原句宣稱「同一 judge」，與上句揭露直接矛盾，故刪去該子句。
        "1d-2 表六說明：刪去矛盾之「與同一 judge」",
        "本表以 CRSCORE++ 評分法在 comprehensiveness",
        "由於四欄均採同一套 CRSCORE++ 提示詞與同一 judge，表六適合",
        "由於四欄均採同一套 CRSCORE++ 提示詞，表六適合",
    ),
    (
        "2 SARIF 重複 gloss（§3.8）",
        "SARIF 匯出（--sarif-out）",
        "以 SARIF（Static Analysis Results Interchange Format，"
        "靜態分析結果交換格式）2.1.0 輸出",
        "以 SARIF 2.1.0 輸出",
    ),
    (
        "3 MCP 重複展開（§3.5）",
        "直覺先行：審查若能在開發者寫程式之當下就給回饋",
        "其一為 Model Context Protocol（MCP）整合層",
        "其一為 MCP 整合層",
    ),
    (
        "3b MCP 第三處 gloss（§3.7.2）",
        "本框架除 CI 觸發路徑外，另以 prthinker mcp 子指令啟動",
        "stdio MCP server（Model Context Protocol，"
        "LLM client 與外部工具之間之 JSON-RPC 協定）",
        "stdio MCP server",
    ),
]


def find_paragraph(doc, anchor):
    hits = [p for p in doc.paragraphs if anchor in p.text]
    if len(hits) != 1:
        raise SystemExit(f"ANCHOR NOT UNIQUE ({len(hits)} hits): {anchor}")
    return hits[0]


def rebuild(p, new_text):
    """Replace all runs with a single run carrying the original rPr plus
    the four explicit rFonts slots."""
    src_rpr = None
    for r in p.runs:
        if r.text.strip():
            src_rpr = r._element.find(qn("w:rPr"))
            break
    for r in list(p.runs):
        r._element.getparent().remove(r._element)
    run = p.add_run(new_text)
    if src_rpr is not None:
        run._element.insert(0, copy.deepcopy(src_rpr))
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    for slot in ("w:ascii", "w:hAnsi", "w:cs"):
        rfonts.set(qn(slot), "Times New Roman")
    rfonts.set(qn("w:eastAsia"), "標楷體")


doc = Document(SRC)
for label, anchor, old, new in EDITS:
    p = find_paragraph(doc, anchor)
    n = p.text.count(old)
    if n != 1:
        raise SystemExit(f"OLD TEXT NOT UNIQUE IN PARAGRAPH ({n} hits): {label}")
    rebuild(p, p.text.replace(old, new))
    # `old` may be a prefix of `new` (pure insertion), so verify on `new`
    # plus the expected total occurrence count instead of `old`'s absence.
    if p.text.count(new) != 1:
        raise SystemExit(f"REPLACEMENT VERIFY FAILED: {label}")
    print(f"OK {label}")

doc.save(DST)
print(f"SAVED {DST.name}")
