"""Apply the TCSE drop-in paragraphs from paper_inserts.md into TCSE_v2.3.docx,
saving a new TCSE_v2.4.docx (original untouched). TCSE inserts are plain body
paragraphs (no new headings), so each blockquote paragraph becomes one Normal
10pt paragraph, inheriting the anchor run's font via a copied rPr."""
import copy
import re
import sys
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph

sys.stdout.reconfigure(encoding="utf-8")
base = Path(__file__).parent
MD = (base / "paper_inserts.md").read_text(encoding="utf-8").split("\n")

# CJK ranges: punctuation, quotes, unified ideographs, fullwidth forms.
_CJK = r"‘-‟　-〿㐀-鿿＀-￯"
_INLINE = re.compile(r"\*\*(.+?)\*\*|``([^`]+?)``|`([^`]+?)`")


def _block_start(header_prefix):
    """Index of the INSERT block heading line in MD."""
    return next(i for i, line in enumerate(MD) if line.startswith(header_prefix))


def _block_lines(header_prefix):
    """Yield the content-area lines of the INSERT block at header_prefix."""
    in_content = False
    for line in MD[_block_start(header_prefix) + 1:]:
        if line.startswith("### ") or line.rstrip() == "---":
            return
        if line.startswith("**內容**") or line.startswith("**標題與內容**"):
            in_content = True
            continue
        if in_content:
            yield line


def _group_paragraphs(lines):
    """Group blockquote lines into paragraphs; blank lines separate them."""
    groups, cur = [], []
    for line in lines:
        if not line.startswith(">"):
            if line.strip() == "" and cur:   # blank line also separates paragraphs
                groups.append(cur)
                cur = []
            continue
        content = line[1:].strip()
        if content:
            cur.append(content)
        elif cur:
            groups.append(cur)
            cur = []
    if cur:
        groups.append(cur)
    return groups


def extract_paras(header_prefix):
    """Return the blockquote paragraphs of an INSERT block as merged strings."""
    out = []
    for lines in _group_paragraphs(_block_lines(header_prefix)):
        text = " ".join(lines).replace("\\ ", "")           # RST zero-width escape
        text = re.sub(rf"(?<=[{_CJK}])\s+(?=[{_CJK}])", "", text)  # drop CJK-internal spaces
        out.append(text)
    return out


def parse_inline(text):
    """Split text into (segment, is_bold) runs, honouring **bold** / ``code`` / `code`."""
    out, pos = [], 0
    for m in _INLINE.finditer(text):
        if m.start() > pos:
            out.append((text[pos:m.start()], False))
        if m.group(1) is not None:
            out.append((m.group(1), True))
        else:
            out.append((m.group(2) or m.group(3), False))
        pos = m.end()
    if pos < len(text):
        out.append((text[pos:], False))
    return out or [(text, False)]


def find_anchor(doc, sub):
    for p in doc.paragraphs:
        if sub in p.text:
            return p
    raise SystemExit(f"ANCHOR NOT FOUND: {sub}")


def template_rpr(anchor):
    for r in anchor.runs:
        if r.text.strip():
            return r._element.find(qn("w:rPr"))
    return None


def insert_after(doc, anchor, texts):
    rpr = template_rpr(anchor)
    ppr = anchor._p.find(qn("w:pPr"))  # copy the anchor's paragraph format (indent, spacing, style)
    cursor = anchor
    for text in texts:
        new_p = OxmlElement("w:p")
        cursor._p.addnext(new_p)
        if ppr is not None:
            new_p.append(copy.deepcopy(ppr))
        para = Paragraph(new_p, anchor._parent)
        if ppr is None:
            para.style = doc.styles["Normal"]
        for seg, bold in parse_inline(text):
            run = para.add_run(seg)
            if rpr is not None:
                run._element.insert(0, copy.deepcopy(rpr))
            if bold:
                run.bold = True
        cursor = para
    return cursor


doc = Document(base / "TCSE_v2.3.docx")
plan = [
    ("### 1.1 ", "並促成人類審查者與智慧系統之互補合作"),
    ("### 1.2 ", "推論層：系統將構建完成之提示詞"),
    ("### 1.3 ", "本研究仍有若干限制與發展方向"),
]
for header, anchor_sub in plan:
    paras = extract_paras(header)
    anchor = find_anchor(doc, anchor_sub)
    insert_after(doc, anchor, paras)
    print(f"OK {header.strip()}: +{len(paras)} paras after «{anchor_sub[:18]}…»")

def trim_sub(doc, find_sub, old_part, new_part):
    """Replace a redundant substring inside an EXISTING body paragraph (trims the
    original TCSE prose). Rebuilds the paragraph as one run carrying the first
    run's formatting — fine for the single-format 10pt body paragraphs here."""
    p = find_anchor(doc, find_sub)
    if old_part not in p.text:
        raise SystemExit(f"TRIM part not found near «{find_sub[:18]}…»")
    new_full = p.text.replace(old_part, new_part)
    rpr = p.runs[0]._element.find(qn("w:rPr")) if p.runs else None
    for r in list(p.runs):
        r._element.getparent().remove(r._element)
    run = p.add_run(new_full)
    if rpr is not None:
        run._element.insert(0, copy.deepcopy(rpr))


# §2 文獻探討 — drop the AI-cliché summary sentence (綜上所述…), keep the 然而… pivot.
trim_sub(
    doc, "傳統工具在規則僵化與語義理解",
    "綜上所述，結合 LLM、靜態分析工具與各類模型優化技術，將有助於建立更高效、智慧化且可擴展之程式碼審查機制。",
    "",
)
# §3 離線訓練層 — merge the two overlapping "注入少量參數微調" descriptions.
trim_sub(
    doc, "本研究之離線訓練流程結合知識蒸餾與 LoRA 微調",
    "僅針對基座模型注入少量可訓練參數進行微調，使學生模型保留原有通用能力之同時，習得教師模型之審查與推理風格。完成資料準備後，系統載入學生模型並注入 LoRA 適配器（adapter）與量化設定，使模型僅須更新少量低秩矩陣參數即可進行微調，有效降低顯存需求與訓練時間。",
    "載入學生模型並注入 LoRA 適配器（adapter）與量化設定，僅針對基座模型更新少量低秩矩陣參數進行微調，使學生模型於保留原有通用能力之同時習得教師模型之審查與推理風格，並有效降低顯存需求與訓練時間。",
)

doc.save(base / "TCSE_v2.4.docx")
print("SAVED TCSE_v2.4.docx (含 §2 / §3 原文精簡)")
